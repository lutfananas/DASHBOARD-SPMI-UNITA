#!/usr/bin/env python3
"""
LAPORAN EVALUASI DIRI (LED) - APT 4.1
Universitas Tulungagung
Berdasarkan SPMI UNITA 2025 + Hasil AMI 2024-2025
Target: 100-150 halaman
Style: Professional Academic Report
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# COLORS (Dark Navy Theme - sesuai template)
# ============================================================
NAVY = RGBColor(0x1A, 0x2B, 0x5C)        # Dark navy
NAVY_LIGHT = RGBColor(0x2D, 0x4A, 0x8C)   # Lighter navy
BLUE = RGBColor(0x1E, 0x88, 0xE5)          # Bright blue
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF0)
EMERALD = RGBColor(0x10, 0x98, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
ROSE = RGBColor(0xF4, 0x3F, 0x5E)

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color="000000", size="4"):
    """Set cell borders"""
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def add_heading_custom(doc, text, level=1, color=NAVY):
    """Add custom styled heading"""
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="1A2B5C"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        run.font.size = Pt(12)
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, justify=True, indent=True, size=11):
    """Add body paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    """Add bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 + level * 0.75)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_numbered(doc, text):
    """Add numbered item"""
    p = doc.add_paragraph(style='List Number')
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table_data(doc, headers, rows, col_widths=None, header_color="1A2B5C"):
    """Add a styled data table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, header_color)
        set_cell_borders(cell, "1A2B5C", "4")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_data)
            set_cell_borders(cell, "CCCCCC", "2")
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F5F7FA")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = DARK_GRAY
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    
    # Set row heights for header
    table.rows[0].height = Cm(0.8)
    
    return table

def add_info_box(doc, title, content, color="1A2B5C"):
    """Add an info/callout box using a single-cell table"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F0F4FF")
    set_cell_borders(cell, color, "4")
    
    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    
    # Content
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK_GRAY
    run2.font.name = 'Calibri'
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.line_spacing = Pt(14)
    
    return table

def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(6)

# ============================================================
# DOCUMENT SETUP
# ============================================================

doc = Document()

# Page setup - A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK_GRAY
style.paragraph_format.line_spacing = Pt(18)
style.paragraph_format.space_after = Pt(6)

# ============================================================
# COVER PAGE
# ============================================================

# Top spacer
add_spacer(doc, 6)

# Logo placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[LOGO UNIVERSITAS TULUNGAGUNG]")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True

add_spacer(doc, 3)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LAPORAN EVALUASI DIRI")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AKREDITASI PERGURUAN TINGGI (APT 4.1)")
run.font.size = Pt(16)
run.font.color.rgb = NAVY_LIGHT
run.font.name = 'Calibri'
run.bold = True
p.paragraph_format.space_after = Pt(20)

# Divider line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="1A2B5C"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)
add_spacer(doc, 2)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UNIVERSITAS TULUNGAGUNG")
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tahun Akademik 2024/2025")
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(30)

add_spacer(doc, 6)

# Info box at bottom
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Disusun oleh", "Pusat Penjaminan Mutu (PPM) Universitas Tulungagung"),
    ("Ketua PPM", "Dr. Anang Sugeng Cahyono, S.AP., S.H., M.Si."),
    ("Rektor", "Dr. Muharsono, M.Si."),
    ("Tanggal Penetapan", "1 September 2025"),
]
for i, (label, value) in enumerate(info_data):
    cell_label = table.rows[i].cells[0]
    cell_value = table.rows[i].cells[1]
    cell_label.text = label
    cell_value.text = value
    set_cell_shading(cell_label, "1A2B5C")
    set_cell_shading(cell_value, "F0F4FF")
    set_cell_borders(cell_label, "1A2B5C", "2")
    set_cell_borders(cell_value, "1A2B5C", "2")
    for p in cell_label.paragraphs:
        for run in p.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in cell_value.paragraphs:
        for run in p.runs:
            run.font.color.rgb = NAVY
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_label.width = Cm(5)
    cell_value.width = Cm(10)

add_page_break(doc)

# ============================================================
# KATA PENGANTAR
# ============================================================

add_heading_custom(doc, "KATA PENGANTAR", level=1)

add_body(doc, "Puji syukur senantiasa kami panjatkan ke hadirat Tuhan Yang Maha Esa atas segala rahmat, hidayah, dan karunia-Nya, sehingga Pusat Penjaminan Mutu Universitas Tulungagung dapat menyelesaikan penyusunan Laporan Evaluasi Diri (LED) untuk Akreditasi Perguruan Tinggi (APT 4.1) Tahun Akademik 2024/2025. Penyusunan laporan ini merupakan wujud nyata komitmen Universitas Tulungagung dalam menjalankan amanat Peraturan Menteri Pendidikan Tinggi, Sains, dan Teknologi Nomor 39 Tahun 2025 tentang Penjaminan Mutu Pendidikan Tinggi.")

add_body(doc, "Laporan Evaluasi Diri ini disusun berdasarkan hasil Audit Mutu Internal (AMI) yang dilaksanakan pada Tahun Akademik 2024/2025, mencakup 13 dokumen AMI dengan total 212 indikator penilaian yang terdiri dari 194 indikator tercapai (91,5%) dan 18 indikator belum tercapai (8,5%). Selain itu, laporan ini juga mengintegrasikan data dari empat dokumen mutu SPMI Universitas Tulungagung Tahun 2025, yaitu Kebijakan Mutu, Manual Mutu, Standar Mutu, dan Formulir Mutu, yang seluruhnya telah diselaraskan dengan Permen 39 Tahun 2025.")

add_body(doc, "Proses penyusunan LED ini mengikuti prinsip triangulasi data sebagaimana diatur dalam Pasal 66 Permen 39 Tahun 2025, dengan mengintegrasikan tiga sumber data: (1) Pangkalan Data Pendidikan Tinggi (PD Dikti) sebagai sumber data primer; (2) hasil evaluasi internal melalui AMI, evaluasi diri, monitoring rutin, dan Rapat Tinjauan Manajemen (RTM); serta (3) umpan balik dari stakeholder eksternal termasuk hasil akreditasi BAN-PT/LAM, umpan balik pengguna lulusan, dan masukan dari dunia usaha, dunia industri, dan dunia kerja (DUDI).")

add_body(doc, "Laporan ini mengevaluasi 9 kriteria sesuai Instrumen Akreditasi Perguruan Tinggi (IAPT) 4.1, mencakup: (1) Visi, Misi, Tujuan, dan Strategi; (2) Tata Pamong, Tata Kelola, dan Kerja Sama; (3) Mahasiswa; (4) Sumber Daya Manusia; (5) Keuangan, Sarana, dan Prasarana; (6) Pendidikan; (7) Penelitian; (8) Pengabdian kepada Masyarakat; serta (9) Luaran dan Capaian Tridharma. Setiap kriteria dianalisis berdasarkan deskripsi kondisi saat ini, evaluasi terhadap standar yang ditetapkan, identifikasi kekuatan dan kelemahan, serta rekomendasi perbaikan berkelanjutan.")

add_body(doc, "Kami menyampaikan terima kasih yang tulus kepada Rektor, Wakil Rektor, Dekan, Ketua Program Studi, Ketua LPPM, Ketua PPM, dan seluruh sivitas akademika Universitas Tulungagung yang telah memberikan kontribusi, masukan, dan dukungan dalam proses penyusunan laporan ini. Kritik, saran, dan masukan dari berbagai pihak senantiasa kami harapkan demi penyempurnaan laporan ini di masa mendatang.")

add_spacer(doc, 3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Tulungagung, 1 September 2025\nPusat Penjaminan Mutu\nKetua,\n\n\n\nAnang Sugeng Cahyono, S.AP., S.H., M.Si.")
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'

add_page_break(doc)

# ============================================================
# DAFTAR ISI (placeholder)
# ============================================================

add_heading_custom(doc, "DAFTAR ISI", level=1)
add_spacer(doc, 2)

isi = [
    ("KATA PENGANTAR", "i"),
    ("DAFTAR ISI", "iii"),
    ("DAFTAR TABEL", "iv"),
    ("DAFTAR GAMBAR", "v"),
    ("DAFTAR SINGKATAN", "vi"),
    ("", ""),
    ("BAB I PENDAHULUAN", "1"),
    ("  1.1 Latar Belakang", "1"),
    ("  1.2 Dasar Hukum", "3"),
    ("  1.3 Tujuan Evaluasi Diri", "4"),
    ("  1.4 Ruang Lingkup Evaluasi", "5"),
    ("  1.5 Metodologi Evaluasi Diri", "6"),
    ("  1.6 Sistem Penjaminan Mutu Internal (SPMI)", "8"),
    ("", ""),
    ("BAB II PROFIL UNIVERSITAS TULUNGAGUNG", "11"),
    ("  2.1 Sejarah dan Perkembangan", "11"),
    ("  2.2 Visi, Misi, Tujuan, dan Strategi", "13"),
    ("  2.3 Struktur Organisasi", "15"),
    ("  2.4 Program Studi dan Fakultas", "16"),
    ("  2.5 Civitas Akademika", "18"),
    ("", ""),
    ("BAB III KRITERIA 1: VISI, MISI, TUJUAN, DAN STRATEGI", "20"),
    ("  3.1 Deskripsi Kriteria", "20"),
    ("  3.2 Analisis dan Evaluasi", "23"),
    ("  3.3 Hasil AMI 2024/2025", "27"),
    ("  3.4 Kekuatan dan Kelemahan", "29"),
    ("  3.5 Rekomendasi", "31"),
    ("", ""),
    ("BAB IV KRITERIA 2: TATA PAMONG, TATA KELOLA, DAN KERJA SAMA", "33"),
    ("  4.1 Deskripsi Kriteria", "33"),
    ("  4.2 Analisis dan Evaluasi", "36"),
    ("  4.3 Hasil AMI 2024/2025", "40"),
    ("  4.4 Kekuatan dan Kelemahan", "42"),
    ("  4.5 Rekomendasi", "44"),
    ("", ""),
    ("BAB V KRITERIA 3: MAHASISWA", "46"),
    ("  5.1 Deskripsi Kriteria", "46"),
    ("  5.2 Analisis dan Evaluasi", "49"),
    ("  5.3 Hasil AMI 2024/2025", "53"),
    ("  5.4 Kekuatan dan Kelemahan", "55"),
    ("  5.5 Rekomendasi", "57"),
    ("", ""),
    ("BAB VI KRITERIA 4: SUMBER DAYA MANUSIA", "59"),
    ("  6.1 Deskripsi Kriteria", "59"),
    ("  6.2 Analisis dan Evaluasi", "62"),
    ("  6.3 Hasil AMI 2024/2025", "66"),
    ("  6.4 Kekuatan dan Kelemahan", "68"),
    ("  6.5 Rekomendasi", "70"),
    ("", ""),
    ("BAB VII KRITERIA 5: KEUANGAN, SARANA, DAN PRASARANA", "72"),
    ("  7.1 Deskripsi Kriteria", "72"),
    ("  7.2 Analisis dan Evaluasi", "76"),
    ("  7.3 Hasil AMI 2024/2025", "80"),
    ("  7.4 Kekuatan dan Kelemahan", "83"),
    ("  7.5 Rekomendasi", "85"),
    ("", ""),
    ("BAB VIII KRITERIA 6: PENDIDIKAN", "87"),
    ("  8.1 Deskripsi Kriteria", "87"),
    ("  8.2 Analisis dan Evaluasi", "91"),
    ("  8.3 Hasil AMI 2024/2025", "95"),
    ("  8.4 Kekuatan dan Kelemahan", "99"),
    ("  8.5 Rekomendasi", "101"),
    ("", ""),
    ("BAB IX KRITERIA 7: PENELITIAN", "103"),
    ("  9.1 Deskripsi Kriteria", "103"),
    ("  9.2 Analisis dan Evaluasi", "106"),
    ("  9.3 Hasil AMI 2024/2025", "110"),
    ("  9.4 Kekuatan dan Kelemahan", "112"),
    ("  9.5 Rekomendasi", "114"),
    ("", ""),
    ("BAB X KRITERIA 8: PENGABDIAN KEPADA MASYARAKAT", "116"),
    ("  10.1 Deskripsi Kriteria", "116"),
    ("  10.2 Analisis dan Evaluasi", "119"),
    ("  10.3 Hasil AMI 2024/2025", "123"),
    ("  10.4 Kekuatan dan Kelemahan", "125"),
    ("  10.5 Rekomendasi", "127"),
    ("", ""),
    ("BAB XI KRITERIA 9: LUARAN DAN CAPAIAN TRIDHARMA", "129"),
    ("  11.1 Deskripsi Kriteria", "129"),
    ("  11.2 Analisis dan Evaluasi", "132"),
    ("  11.3 Hasil AMI 2024/2025", "135"),
    ("  11.4 Kekuatan dan Kelemahan", "137"),
    ("  11.5 Rekomendasi", "139"),
    ("", ""),
    ("BAB XII PENUTUP", "141"),
    ("  12.1 Ringkasan Hasil Evaluasi Diri", "141"),
    ("  12.2 Rencana Tindak Lanjut", "143"),
    ("  12.3 Penutup", "145"),
    ("", ""),
    ("LAMPIRAN", "147"),
    ("  Lampiran 1: Daftar Dokumen AMI 2024/2025", "147"),
    ("  Lampiran 2: Rekapitulasi 33 Standar SPMI", "149"),
    ("  Lampiran 3: 8 IKU dan Target Capaian", "151"),
]

for item, page in isi:
    p = doc.add_paragraph()
    if item == "":
        p.paragraph_format.space_after = Pt(2)
        continue
    run = p.add_run(f"{item}")
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GRAY
    if item.startswith("BAB") or item.startswith("LAMPIRAN"):
        run.bold = True
        run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)
    
    # Tab to page number
    run2 = p.add_run(f"\t{page}")
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    run2.font.color.rgb = GRAY

add_page_break(doc)

# ============================================================
# BAB I - PENDAHULUAN
# ============================================================

add_heading_custom(doc, "BAB I\nPENDAHULUAN", level=1)

add_heading_custom(doc, "1.1 Latar Belakang", level=2)

add_body(doc, "Universitas Tulungagung sebagai institusi pendidikan tinggi memiliki tanggung jawab strategis dalam mencerdaskan kehidupan bangsa, memajukan ilmu pengetahuan dan teknologi, serta menghasilkan sumber daya manusia unggul yang mampu berkompetisi pada skala nasional dan internasional. Dalam menjalankan tanggung jawab tersebut, Universitas Tulungagung wajib memastikan bahwa seluruh penyelenggaraan Tridharma Perguruan Tinggi—yang meliputi pendidikan, penelitian, dan pengabdian kepada masyarakat—dilaksanakan sesuai dengan standar mutu yang telah ditetapkan, baik standar nasional maupun standar yang ditetapkan oleh perguruan tinggi itu sendiri.")

add_body(doc, "Berlakunya Peraturan Menteri Pendidikan Tinggi, Sains, dan Teknologi Nomor 39 Tahun 2025 tentang Penjaminan Mutu Pendidikan Tinggi sejak tanggal 2 September 2025 membawa konsekuensi yuridis dan operasional yang signifikan bagi setiap perguruan tinggi. Peraturan ini menjadi landasan regulasi terkini yang mengatur kerangka Sistem Penjaminan Mutu Pendidikan Tinggi (SPM Dikti), yang mencakup Sistem Penjaminan Mutu Internal (SPMI) dan Sistem Penjaminan Mutu Eksternal (SPME) melalui Akreditasi. Universitas Tulungagung merespons regulasi ini dengan menyusun dan memperbarui seluruh dokumen SPMI, serta melaksanakan Audit Mutu Internal (AMI) secara komprehensif pada Tahun Akademik 2024/2025.")

add_body(doc, "Laporan Evaluasi Diri (LED) ini disusun sebagai bagian dari proses akreditasi perguruan tinggi sesuai Instrumen Akreditasi Perguruan Tinggi (IAPT) 4.1. Laporan ini merupakan hasil refleksi dan analisis kritis terhadap kinerja Universitas Tulungagung dalam menyelenggarakan Tridharma, dengan mengacu pada hasil AMI 2024/2025 yang mencakup 13 dokumen audit dengan 212 indikator penilaian. Selain itu, laporan ini juga mengintegrasikan data dari empat dokumen mutu SPMI yang telah diselaraskan dengan Permen 39 Tahun 2025, yaitu Kebijakan Mutu (65 halaman), Manual Mutu (55 halaman), Standar Mutu (312 halaman), dan Formulir Mutu (132 halaman).")

add_body(doc, "Penyusunan LED ini mengikuti prinsip triangulasi data sebagaimana diatur dalam Pasal 66 Permen 39 Tahun 2025, yang mengharuskan penggalian kebenaran informasi melalui penggunaan berbagai sumber data dan sudut pandang yang saling melengkapi. Tiga sumber data yang digunakan adalah: (1) Pangkalan Data Pendidikan Tinggi (PD Dikti) sebagai sumber data primer; (2) hasil evaluasi internal melalui AMI, evaluasi diri, monitoring rutin, dan Rapat Tinjauan Manajemen (RTM); serta (3) umpan balik dari stakeholder eksternal termasuk hasil akreditasi BAN-PT/LAM, umpan balik pengguna lulusan, dan masukan dari DUDI.")

add_body(doc, "Laporan ini mengevaluasi 9 kriteria sesuai IAPT 4.1, yang mencakup seluruh aspek penyelenggaraan perguruan tinggi. Setiap kriteria dianalisis berdasarkan deskripsi kondisi saat ini, evaluasi terhadap standar yang ditetapkan, identifikasi kekuatan dan kelemahan, serta rekomendasi perbaikan berkelanjutan. Hasil evaluasi diri ini akan menjadi dasar bagi penyusunan Rencana Tindak Lanjut (RTL) dan strategi peningkatan mutu Universitas Tulungagung ke depan.")

add_heading_custom(doc, "1.2 Dasar Hukum", level=2)

add_body(doc, "Penyusunan Laporan Evaluasi Diri ini didasarkan pada peraturan perundang-undangan yang berlaku, yang menjadi landasan yuridis dan operasional bagi seluruh kegiatan penjaminan mutu di lingkungan Universitas Tulungagung. Dasar hukum ini merupakan kerangka regulasi terkini yang berlaku efektif sejak tanggal diundangkannya Peraturan Menteri Pendidikan Tinggi, Sains, dan Teknologi Nomor 39 Tahun 2025.")

dasar_hukum = [
    "1. Undang-Undang Republik Indonesia Nomor 20 Tahun 2003 tentang Sistem Pendidikan Nasional;",
    "2. Undang-Undang Republik Indonesia Nomor 12 Tahun 2012 tentang Pendidikan Tinggi;",
    "3. Peraturan Pemerintah Republik Indonesia Nomor 4 Tahun 2014 tentang Penyelenggaraan Pendidikan Tinggi dan Pengelolaan Perguruan Tinggi;",
    "4. Peraturan Menteri Pendidikan Tinggi, Sains, dan Teknologi Nomor 39 Tahun 2025 tentang Penjaminan Mutu Pendidikan Tinggi;",
    "5. Peraturan Badan Akreditasi Nasional Perguruan Tinggi (BAN-PT) yang berlaku;",
    "6. Instrumen Akreditasi Perguruan Tinggi (IAPT) 4.1;",
    "7. Statuta Universitas Tulungagung;",
    "8. Rencana Strategis (Renstra) Universitas Tulungagung Tahun 2023-2027;",
    "9. Keputusan Rektor Universitas Tulungagung Nomor A/002.I/KEP/UNITA/I/2025 tentang Pernyataan dan Kebijakan Mutu Universitas Tulungagung.",
]
for item in dasar_hukum:
    add_body(doc, item, indent=False)

add_heading_custom(doc, "1.3 Tujuan Evaluasi Diri", level=2)

add_body(doc, "Laporan Evaluasi Diri ini disusun dengan tujuan untuk: (1) memberikan gambaran objektif dan komprehensif tentang kondisi Universitas Tulungagung dalam penyelenggaraan Tridharma Perguruan Tinggi berdasarkan 9 kriteria IAPT 4.1; (2) mengevaluasi ketercapaian 33 Standar SPMI yang telah ditetapkan berdasarkan hasil AMI 2024/2025; (3) mengidentifikasi kekuatan, kelemahan, peluang, dan tantangan dalam setiap kriteria; (4) menyusun rekomendasi dan rencana tindak lanjut untuk peningkatan mutu berkelanjutan; serta (5) memenuhi persyaratan akreditasi perguruan tinggi sesuai regulasi yang berlaku.")

add_heading_custom(doc, "1.4 Ruang Lingkup Evaluasi", level=2)

add_body(doc, "Ruang lingkup evaluasi diri ini mencakup seluruh aspek penyelenggaraan Universitas Tulungagung, yang dikelompokkan ke dalam 9 kriteria sesuai IAPT 4.1: (1) Visi, Misi, Tujuan, dan Strategi (VMTS); (2) Tata Pamong, Tata Kelola, dan Kerja Sama; (3) Mahasiswa; (4) Sumber Daya Manusia; (5) Keuangan, Sarana, dan Prasarana; (6) Pendidikan; (7) Penelitian; (8) Pengabdian kepada Masyarakat; serta (9) Luaran dan Capaian Tridharma. Setiap kriteria dievaluasi berdasarkan data dari AMI 2024/2025, dokumen SPMI, PD Dikti, dan umpan balik stakeholder.")

add_heading_custom(doc, "1.5 Metodologi Evaluasi Diri", level=2)

add_body(doc, "Metodologi evaluasi diri menggunakan pendekatan kuantitatif dan kualitatif dengan prinsip triangulasi data. Data dikumpulkan dari tiga sumber: (1) PD Dikti sebagai sumber data primer yang mencakup data mahasiswa, dosen, kurikulum, lulusan, dan Tridharma; (2) hasil evaluasi internal yang meliputi 13 dokumen AMI 2024/2025 dengan 212 indikator, laporan RTM, dan laporan monitoring rutin; serta (3) data stakeholder eksternal yang meliputi hasil akreditasi BAN-PT/LAM, tracer study, umpan balik DUDI, dan pengaduan masyarakat.")

add_body(doc, "Proses evaluasi diri dilaksanakan melalui tahapan: (1) pengumpulan data dari berbagai sumber; (2) analisis data berdasarkan standar IAPT 4.1 dan 33 Standar SPMI; (3) identifikasi kekuatan dan kelemahan; (4) penyusunan rekomendasi perbaikan; dan (5) validasi oleh Senat Universitas. Hasil evaluasi diri ini selanjutnya dibahas dalam Rapat Tinjauan Manajemen (RTM) dan ditindaklanjuti melalui Rencana Tindak Lanjut (RTL).")

add_heading_custom(doc, "1.6 Sistem Penjaminan Mutu Internal (SPMI)", level=2)

add_body(doc, "Universitas Tulungagung telah menerapkan Sistem Penjaminan Mutu Internal (SPMI) yang terstruktur dan terdokumentasi, sesuai dengan amanat Pasal 67 Permen 39 Tahun 2025. SPMI Universitas Tulungagung dikelola melalui siklus PPEPP (Penetapan, Pelaksanaan, Evaluasi, Pengendalian, dan Peningkatan) sebagaimana diatur dalam Pasal 68. Sistem ini didukung oleh empat dokumen mutu utama yang telah diselaraskan dengan Permen 39 Tahun 2025.")

add_info_box(doc, "Dokumen Mutu SPMI UNITA 2025",
    "1. Kebijakan Mutu SPMI (Kode: SPMI/PPM/DM/KBJ/2025) — 65 halaman\n"
    "2. Manual Mutu SPMI (Kode: SPMI/PPM/DM/MNL/2025) — 55 halaman\n"
    "3. Standar Mutu SPMI (Kode: SPMI/PPM/DM/STD/2025) — 312 halaman\n"
    "4. Formulir Mutu SPMI (Kode: SPMI/PPM/DM/FORM/2025) — 132 halaman\n\n"
    "Total: 33 Standar SPMI (8 Pendidikan + 8 Penelitian + 8 PkM + 9 Tambahan)\n"
    "Total formulir: 38 formulir pendukung PPEPP"
)

add_spacer(doc, 2)

add_body(doc, "Selain itu, Universitas Tulungagung telah melaksanakan Audit Mutu Internal (AMI) pada Tahun Akademik 2024/2025 yang mencakup 13 dokumen audit dengan total 212 indikator penilaian. Hasil AMI menunjukkan bahwa 194 indikator (91,5%) berstatus Tercapai dan 18 indikator (8,5%) berstatus Belum Tercapai. Hasil AMI ini menjadi data utama dalam penyusunan Laporan Evaluasi Diri ini.")

# Table: Ringkasan Hasil AMI 2024/2025
add_spacer(doc, 2)
p = doc.add_paragraph()
run = p.add_run("Tabel 1.1 Ringkasan Hasil AMI 2024/2025")
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_table_data(doc,
    ["No", "Dokumen AMI", "Kode", "Indikator", "Tercapai", "Belum"],
    [
        ["1", "AMI Pendidikan", "MONEV/PEND/2025", "19", "19", "0"],
        ["2", "AMI Penelitian", "MONEV/PEN/2025", "20", "20", "0"],
        ["3", "AMI Pengabdian Masyarakat", "MONEV/PKM/2025", "20", "20", "0"],
        ["4", "AMI VMTS", "MONEV/VMTS/2025", "15", "10", "5"],
        ["5", "AMI TPTK", "MONEV/TPTK/2025", "15", "12", "3"],
        ["6", "AMI Mahasiswa", "MONEV/MHS/2025", "14", "14", "0"],
        ["7", "AMI SDM", "MONEV/SDM/2025", "18", "16", "2"],
        ["8", "AMI Kesejahteraan", "MONEV/KSJH/2025", "19", "19", "0"],
        ["9", "AMI Keuangan", "MONEV/KEU/2025", "14", "14", "0"],
        ["10", "AMI Sarana Prasarana", "MONEV/SARPRAS/2025", "14", "14", "0"],
        ["11", "AMI MBKM", "MONEV/MBKM/2025", "4", "4", "0"],
        ["12", "AMI Pelayanan Mahasiswa", "MONEV/PLYN/2025", "20", "13", "7"],
        ["13", "AMI Proses Pembelajaran", "MONEV/SPB/2025", "20", "20", "0"],
        ["", "TOTAL", "", "212", "194", "18"],
    ],
    col_widths=[1, 4, 3, 2, 2, 2]
)

add_page_break(doc)

# ============================================================
# BAB II - PROFIL UNIVERSITAS TULUNGAGUNG
# ============================================================

add_heading_custom(doc, "BAB II\nPROFIL UNIVERSITAS TULUNGAGUNG", level=1)

add_heading_custom(doc, "2.1 Sejarah dan Perkembangan", level=2)

add_body(doc, "Universitas Tulungagung (UNITA) merupakan salah satu perguruan tinggi swasta di Kabupaten Tulungagung, Jawa Timur, yang berdiri dengan komitmen untuk menyelenggarakan pendidikan tinggi yang berkualitas dan mampu berkompetisi pada skala nasional dan internasional. Sejak berdirinya, Universitas Tulungagung terus berkembang baik dari sisi jumlah program studi, jumlah mahasiswa, kualitas dosen, maupun sarana dan prasarana pendukung kegiatan akademik.")

add_body(doc, "Pada tahun 2011, UNITA mulai menerapkan Sistem Penjaminan Mutu Internal (SPMI) secara formal, dengan mengacu pada standar akreditasi nasional dari BAN-PT. Sejak saat itu, SPMI di UNITA terus dikembangkan dan disempurnakan mengikuti perkembangan regulasi penjaminan mutu pendidikan tinggi di Indonesia. Pada tahun 2025, sejalan dengan berlakunya Permen 39 Tahun 2025, Universitas Tulungagung melakukan penyesuaian menyeluruh terhadap seluruh dokumen SPMI agar selaras dengan kerangka regulasi terbaru.")

add_body(doc, "Saat ini, Universitas Tulungagung memiliki 6 fakultas yang menyelenggarakan berbagai program studi pada jenjang diploma, sarjana, dan pascasarjana. Fakultas-fakultas tersebut meliputi Fakultas Ekonomi (FE), Fakultas Ilmu Sosial dan Ilmu Politik (FISIP), Fakultas Pertanian (FP), Fakultas Hukum (FH), Fakultas Teknik (FT), dan Fakultas Kesehatan (F.KES). Selain itu, UNITA juga memiliki Program Diploma Tiga (D3) Kebidanan.")

add_heading_custom(doc, "2.2 Visi, Misi, Tujuan, dan Strategi", level=2)

add_body(doc, "Visi Universitas Tulungagung adalah:")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\u201cMewujudkan Perguruan Tinggi yang Berkualitas dan Mampu Berkompetisi Berskala Nasional Dan Internasional.\u201d")
run.font.size = Pt(12)
run.font.bold = True
run.italic = True
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)

add_body(doc, "Misi Universitas Tulungagung adalah:")
misi = [
    "Peningkatan penyelenggaraan pendidikan, penelitian, dan pengabdian kepada masyarakat yang berkualitas menuju terbentuknya lulusan mandiri;",
    "Peningkatan kualitas Sumber Daya Manusia yang meliputi tenaga pendidik dan tenaga kependidikan;",
    "Peningkatan sarana dan prasarana akademik yang mendukung proses belajar mengajar.",
]
for item in misi:
    add_numbered(doc, item)

add_body(doc, "Tujuan Universitas Tulungagung adalah:")
tujuan = [
    "Berkembangnya potensi mahasiswa agar menjadi manusia yang beriman dan bertakwa kepada Tuhan Yang Maha Esa dan berakhlak mulia, sehat, berilmu, cakap, kreatif, mandiri, terampil, kompeten, dan berbudaya untuk kepentingan bangsa;",
    "Dihasilkannya lulusan yang menguasai cabang ilmu pengetahuan dan/atau teknologi untuk memenuhi kepentingan nasional dan peningkatan daya saing bangsa;",
    "Dihasilkannya ilmu pengetahuan dan teknologi melalui penelitian yang memperhatikan dan menerapkan nilai humaniora agar bermanfaat bagi kemajuan bangsa, serta kemajuan peradaban dan kesejahteraan umat manusia;",
    "Terwujudnya pengabdian kepada masyarakat berbasis penalaran dan karya penelitian yang bermanfaat dalam memajukan kesejahteraan umum dan mencerdaskan kehidupan bangsa.",
]
for item in tujuan:
    add_numbered(doc, item)

add_heading_custom(doc, "2.3 Struktur Organisasi", level=2)

add_body(doc, "Universitas Tulungagung dipimpin oleh seorang Rektor yang bertanggung jawab atas penyelenggaraan seluruh kegiatan akademik dan non-akademik. Rektor dibantu oleh tiga Wakil Rektor: Wakil Rektor I (Bidang Akademik), Wakil Rektor II (Bidang Umum dan Keuangan), dan Wakil Rektor III (Bidang Kemahasiswaan). Selain itu, terdapat beberapa lembaga dan biro pendukung, antara lain Lembaga Penelitian dan Pengabdian kepada Masyarakat (LPPM), Pusat Penjaminan Mutu (PPM), Biro Administrasi Akademik (BAA), Biro Administrasi Umum (BAU), Biro Administrasi Keuangan (BAKU), dan Biro Administrasi Kemahasiswaan (BAK).")

add_body(doc, "Pusat Penjaminan Mutu (PPM) dipimpin oleh seorang Ketua yang bertanggung jawab langsung kepada Rektor. PPM dibantu oleh Lembaga Penjaminan Mutu (LPM) di tingkat fakultas dan Gugus Penjaminan Mutu (GKM) di tingkat program studi. Struktur organisasi PPM juga mencakup Sekretaris, Kepala Bagian Pengembangan SPMI, Kepala Bagian Audit dan Monevin, serta staf administrasi. Susunan organisasi PPM ditetapkan berdasarkan Surat Keputusan Rektor Nomor A/001.D/KEP/UNITA/I/2025.")

add_heading_custom(doc, "2.4 Program Studi dan Fakultas", level=2)

add_body(doc, "Universitas Tulungagung memiliki 6 fakultas dengan berbagai program studi pada jenjang diploma, sarjana, dan pascasarjana. Berikut adalah daftar fakultas dan program studi di Universitas Tulungagung:")

add_table_data(doc,
    ["No", "Fakultas", "Program Studi", "Jenjang", "Akreditasi"],
    [
        ["1", "Fakultas Ekonomi", "Akuntansi", "S1", "Unggul"],
        ["2", "Fakultas Ekonomi", "Manajemen", "S1", "Unggul"],
        ["3", "Fakultas Ekonomi", "Ekonomi Pembangunan", "S1", "Baik Sekali"],
        ["4", "Fakultas Ekonomi", "Magister Manajemen", "S2", "Baik Sekali"],
        ["5", "Fakultas Hukum", "Ilmu Hukum", "S1", "Unggul"],
        ["6", "Fakultas FISIP", "Ilmu Administrasi Negara", "S1", "Baik Sekali"],
        ["7", "Fakultas FISIP", "Ilmu Komunikasi", "S1", "Baik Sekali"],
        ["8", "Fakultas Pertanian", "Agroteknologi", "S1", "Baik Sekali"],
        ["9", "Fakultas Pertanian", "Agribisnis", "S1", "Baik Sekali"],
        ["10", "Fakultas Teknik", "Teknik Sipil", "S1", "Baik Sekali"],
        ["11", "Fakultas Teknik", "Informatika", "S1", "Baik Sekali"],
        ["12", "Fakultas Kesehatan", "Ilmu Keperawatan", "S1", "Baik Sekali"],
        ["13", "Fakultas Kesehatan", "Kebidanan", "D3", "Baik Sekali"],
    ],
    col_widths=[1, 3, 4, 2, 3]
)

add_heading_custom(doc, "2.5 Civitas Akademika", level=2)

add_body(doc, "Berdasarkan data terbaru Tahun 2025, Universitas Tulungagung memiliki total 93 orang SDM yang terdiri dari dosen dan tenaga kependidikan. Dari jumlah tersebut, sebanyak 78,40% dosen tetap memiliki jabatan fungsional minimal Lektor. Kepemilikan sertifikat pendidik telah mencapai 94,60%, sementara persentase dosen tidak tetap dengan sertifikat kompetensi berada pada kisaran 16,20% dari total dosen, yang masih di bawah batas maksimum 20%.")

add_body(doc, "Jumlah mahasiswa aktif Universitas Tulungagung pada Tahun Akademik 2024/2025 mencapai sekitar 8.000 mahasiswa, dengan rasio dosen-mahasiswa rata-rata 1:35. Universitas Tulungagung juga aktif menerima mahasiswa transfer dan mahasiswa internasional sebagai bagian dari diversifikasi mahasiswa. Tingkat retensi mahasiswa (sampai lulus) berada pada kisaran 75%, dengan target peningkatan menjadi 95% pada tahun 2030.")

add_body(doc, "Lulusan Universitas Tulungagung memiliki tingkat penyerapan di dunia kerja yang tinggi, dengan 86,50% lulusan terserap di dunia kerja dalam waktu kurang dari 6 bulan setelah kelulusan. Rata-rata IPK lulusan adalah 3,48, dengan rata-rata masa studi 3,8 tahun dan rata-rata masa tunggu kerja 3,2 bulan. Sebanyak 18,20% lulusan menjadi wirausahawan, menunjukkan bahwa Universitas Tulungagung tidak hanya menghasilkan lulusan yang siap kerja, tetapi juga mampu menciptakan lapangan kerja.")

add_page_break(doc)

# ============================================================
# HELPER: Generate Kriteria BAB
# ============================================================

def generate_kriteria_bab(doc, bab_num, kriteria_num, kriteria_nama, deskripsi_paragrafs,
                          analisis_paragrafs, ami_data_table, kekuatan_list, kelemahan_list,
                          rekomendasi_list, ami_rata_skor, ami_status, ami_dokumen):
    """Generate a complete kriteria chapter"""
    
    add_heading_custom(doc, f"BAB {bab_num}\nKRITERIA {kriteria_num}: {kriteria_nama.upper()}", level=1)
    
    # Deskripsi
    add_heading_custom(doc, f"{bab_num}.1 Deskripsi Kriteria", level=2)
    for para in deskripsi_paragrafs:
        add_body(doc, para)
    
    # Analisis
    add_heading_custom(doc, f"{bab_num}.2 Analisis dan Evaluasi", level=2)
    for para in analisis_paragrafs:
        add_body(doc, para)
    
    # Hasil AMI
    add_heading_custom(doc, f"{bab_num}.3 Hasil AMI 2024/2025", level=2)
    
    add_info_box(doc, f"Hasil AMI {kriteria_nama}",
        f"Rata-rata Skor: {ami_rata_skor}\n"
        f"Status: {ami_status}\n"
        f"Dokumen AMI: {ami_dokumen}\n"
        f"Sumber: Hasil Audit Mutu Internal Universitas Tulungagung Tahun Akademik 2024/2025"
    )
    
    add_spacer(doc, 2)
    
    # AMI Detail Table
    p = doc.add_paragraph()
    run = p.add_run(f"Tabel {bab_num}.1 Hasil Detail AMI {kriteria_nama} Tahun 2024/2025")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_table_data(doc,
        ["No", "Indikator", "Capaian", "Target", "Status"],
        ami_data_table,
        col_widths=[1, 6, 3, 3, 2]
    )
    
    # Kekuatan dan Kelemahan
    add_heading_custom(doc, f"{bab_num}.4 Kekuatan dan Kelemahan", level=2)
    
    add_heading_custom(doc, "Kekuatan", level=3)
    for item in kekuatan_list:
        add_bullet(doc, item)
    
    add_heading_custom(doc, "Kelemahan", level=3)
    for item in kelemahan_list:
        add_bullet(doc, item)
    
    # Rekomendasi
    add_heading_custom(doc, f"{bab_num}.5 Rekomendasi", level=2)
    for i, item in enumerate(rekomendasi_list):
        add_numbered(doc, item)
    
    add_page_break(doc)


# ============================================================
# BAB III - KRITERIA 1: VMTS
# ============================================================

generate_kriteria_bab(doc, 3, 1, "Visi, Misi, Tujuan, dan Strategi",
    deskripsi_paragrafs=[
        "Kriteria 1 mengevaluasi perumusan, sosialisasi, dan implementasi Visi, Misi, Tujuan, dan Strategi (VMTS) di Universitas Tulungagung. VMTS menjadi landasan filosofis dan strategis bagi seluruh kegiatan perguruan tinggi, sehingga kriteria ini menempati posisi penting dalam evaluasi diri. VMTS Universitas Tulungagung telah ditetapkan melalui proses partisipatif yang melibatkan Senat, dosen, mahasiswa, tenaga kependidikan, alumni, dan pemangku kepentingan eksternal.",
        "Visi Universitas Tulungagung adalah \u201cMewujudkan Perguruan Tinggi yang Berkualitas dan Mampu Berkompetisi Berskala Nasional Dan Internasional.\u201d Visi ini mengandung makna bahwa Universitas Tulungagung berkomitmen untuk terus meningkatkan kualitas penyelenggaraan Tridharma Perguruan Tinggi, sehingga mampu berkompetisi tidak hanya pada level nasional tetapi juga internasional.",
        "VMTS diturunkan ke dalam Rencana Induk Pengembangan (RIP), Rencana Strategis (Renstra) Tahun 2023-2027, dan Rencana Operasional (Renop) tahunan. VMTS juga dipublikasikan kepada masyarakat melalui situs web resmi universitas dan dokumen resmi, sebagai bentuk transparansi dan akuntabilitas.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI VMTS 2024/2025, rata-rata skor keseluruhan adalah 3,54 (skala Likert 1-4) dengan kategori Sangat Baik. Skor ini menunjukkan bahwa pemahaman dan implementasi VMTS di kalangan civitas akademika Universitas Tulungagung berada pada tingkat yang baik. Namun, terdapat beberapa indikator yang masih berada di bawah target 3,50 dan menjadi prioritas perbaikan.",
        "Dari 15 indikator yang dievaluasi, 10 indikator berstatus Tercapai dan 5 indikator berstatus Belum Tercapai. Indikator yang belum tercapai mencakup: sosialisasi VMTS secara berkala (skor 3,42), media dan metode sosialisasi VMTS yang mudah dipahami (skor 3,38), keterlibatan dosen dalam penyusunan/evaluasi VMTS (skor 3,48), keterlibatan dosen dalam evaluasi ketercapaian VMTS (skor 3,36), serta relevansi VMTS dengan kebutuhan dunia kerja (skor 3,48).",
        "Hasil AMI menunjukkan bahwa VMTS telah menjadi acuan dalam pelaksanaan pendidikan (skor 3,62), penelitian (skor 3,56), dan pengabdian kepada masyarakat (skor 3,54). Pimpinan juga menunjukkan komitmen yang kuat dalam implementasi VMTS (skor 3,62). Namun, metode sosialisasi VMTS perlu diperbaiki agar lebih mudah dipahami oleh seluruh sivitas akademika.",
    ],
    ami_data_table=[
        ["1", "Pemahaman VMTS Universitas Tulungagung", "3,55", "≥ 3,50", "Tercapai"],
        ["2", "Pemahaman keterkaitan VMTS dengan program kerja", "3,50", "≥ 3,50", "Tercapai"],
        ["3", "VMTS disosialisasikan secara berkala kepada dosen", "3,42", "≥ 3,50", "Belum"],
        ["4", "Media dan metode sosialisasi VMTS mudah dipahami", "3,38", "≥ 3,50", "Belum"],
        ["5", "VMTS menjadi acuan pelaksanaan pendidikan", "3,62", "≥ 3,50", "Tercapai"],
        ["6", "VMTS menjadi acuan pelaksanaan penelitian", "3,56", "≥ 3,50", "Tercapai"],
        ["7", "VMTS menjadi acuan pelaksanaan pengabdian", "3,54", "≥ 3,50", "Tercapai"],
        ["8", "VMTS terintegrasi dalam kurikulum dan pembelajaran", "3,50", "≥ 3,50", "Tercapai"],
        ["9", "Pimpinan menunjukkan komitmen kuat dalam VMTS", "3,62", "≥ 3,50", "Tercapai"],
        ["10", "Kebijakan pimpinan selaras dengan VMTS", "3,58", "≥ 3,50", "Tercapai"],
        ["11", "Dosen dilibatkan dalam penyusunan/evaluasi VMTS", "3,48", "≥ 3,50", "Belum"],
        ["12", "Dosen dilibatkan dalam evaluasi ketercapaian VMTS", "3,36", "≥ 3,50", "Belum"],
        ["13", "VMTS relevan dengan perkembangan ilmu pengetahuan", "3,50", "≥ 3,50", "Tercapai"],
        ["14", "VMTS relevan dengan kebutuhan dunia kerja", "3,48", "≥ 3,50", "Belum"],
        ["15", "Program universitas mendukung pencapaian VMTS", "3,40", "≥ 3,50", "Belum"],
    ],
    kekuatan_list=[
        "VMTS telah dirumuskan melalui proses partisipatif yang melibatkan seluruh pemangku kepentingan",
        "VMTS telah menjadi acuan dalam pelaksanaan Tridharma Perguruan Tinggi (skor 3,54-3,62)",
        "Pimpinan menunjukkan komitmen yang kuat dalam implementasi VMTS (skor 3,62)",
        "VMTS telah diturunkan ke dalam RIP, Renstra 2023-2027, dan Renop tahunan",
        "Pemahaman civitas akademika terhadap VMTS berada pada kategori Sangat Baik (3,54)",
        "VMTS telah dipublikasikan melalui situs web resmi universitas",
    ],
    kelemahan_list=[
        "Metode sosialisasi VMTS belum optimal (skor 3,38), perlu diversifikasi media",
        "Keterlibatan dosen dalam evaluasi ketercapaian VMTS masih rendah (skor 3,36)",
        "Relevansi VMTS dengan kebutuhan dunia kerja perlu ditingkatkan (skor 3,48)",
        "Program universitas belum sepenuhnya mendukung pencapaian VMTS (skor 3,40)",
        "Review VMTS belum dilakukan secara berkala setiap 5 tahun sebagaimana standar",
    ],
    rekomendasi_list=[
        "Diversifikasi media dan metode sosialisasi VMTS melalui rapat, situs web, media sosial, dan handbook",
        "Melibatkan dosen secara aktif dalam evaluasi ketercapaian VMTS melalui mekanisme feedback",
        "Mengundang DUDI dalam review VMTS untuk memastikan relevansi dengan kebutuhan dunia kerja",
        "Menyelenggarakan workshop penyusunan RIP, Renstra, dan Renop bagi pimpinan unit kerja",
        "Melakukan review VMTS setiap 5 tahun untuk memastikan relevansi dengan dinamika eksternal",
        "Mengembangkan dashboard monitoring ketercapaian VMTS dengan indikator yang terukur",
    ],
    ami_rata_skor="3,54/4 (Sangat Baik)",
    ami_status="10 Tercapai, 5 Belum Tercapai",
    ami_dokumen="AMI VMTS 2024-2025 (Kode: SPMI/PPM/MONEV/VMTS/2025)"
)

# ============================================================
# BAB IV - KRITERIA 2: TATA PAMONG
# ============================================================

generate_kriteria_bab(doc, 4, 2, "Tata Pamong, Tata Kelola, dan Kerja Sama",
    deskripsi_paragrafs=[
        "Kriteria 2 mengevaluasi struktur tata pamong, sistem pengelolaan, sistem penjaminan mutu, dan kerja sama di Universitas Tulungagung. Sesuai Pasal 67 ayat (2) Permen 39 Tahun 2025, Universitas Tulungagung menerapkan tata kelola perguruan tinggi yang baik berdasarkan prinsip akuntabilitas, transparansi, nirlaba, efektivitas, efisiensi, dan peningkatan mutu berkelanjutan, yang saling menilik dan mengimbangi satu terhadap yang lain.",
        "Tata pamong Universitas Tulungagung mencakup struktur organisasi (Rektorat, Fakultas, Lembaga, Biro, UPT), mekanisme pengambilan keputusan (Senat, Rapat Pimpinan, Rapat Unit), sistem pengelolaan (perencanaan, pelaksanaan, monitoring, evaluasi), dan sistem penjaminan mutu (SPMI, AMI, RTM). Setiap elemen tata pamong memiliki tugas, wewenang, dan tanggung jawab yang jelas, dan diatur dalam Statuta dan SOTK.",
        "Universitas Tulungagung juga memiliki mekanisme pengawasan oleh Senat dan Satuan Pengawas Internal (SPI), sebagai bagian dari check-and-balance dalam tata kelola. Senat melakukan pengawasan terhadap kebijakan akademik, sedangkan SPI melakukan audit terhadap pengelolaan keuangan dan operasional. Hasil pengawasan menjadi basis bagi perbaikan berkelanjutan, dan dilaporkan dalam RTM tahunan.",
        "Dalam hal kerja sama, Universitas Tulungagung aktif mengembangkan kerja sama dengan perguruan tinggi lain, DUDI, pemerintah daerah, dan lembaga internasional. Kerja sama mencakup bidang akademik (joint degree, exchange, visiting), penelitian (joint research, hibah bersama), PkM (joint PkM, program bersama), dan kerja sama DUDI (magang, MBKM, penelitian terapan).",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI TPTK 2024/2025, rata-rata skor keseluruhan adalah 3,59 (skala Likert 1-4) dengan kategori Sangat Baik. Skor ini menunjukkan bahwa tata pamong dan tata kelola Universitas Tulungagung berjalan dengan baik. Dari 15 indikator yang dievaluasi, 12 indikator berstatus Tercapai dan 3 indikator berstatus Belum Tercapai.",
        "Indikator yang belum tercapai mencakup: keterlibatan dosen dalam penyusunan dan evaluasi kebijakan (skor 3,48), tata kelola yang mendukung pelaksanaan Tridharma (skor 3,48), serta keterbukaan pimpinan terhadap masukan dosen (skor 3,49). Ketiga indikator ini berada di bawah target 3,50 dan menjadi prioritas perbaikan.",
        "Untuk kerja sama, berdasarkan data AMI MBKM, jumlah kerja sama dengan mitra luar kampus mencapai skor 92,40 (Sangat Baik). Universitas Tulungagung telah memiliki MoU dengan berbagai perguruan tinggi dan DUDI, dengan implementasi yang terus ditingkatkan. Saat ini terdapat sekitar 40 MoU aktif dan 60 MoA aktif, dengan 10 kerja sama internasional.",
    ],
    ami_data_table=[
        ["1", "Pemahaman struktur tata pamong dan tata kelola", "3,58", "≥ 3,50", "Tercapai"],
        ["2", "Kebijakan akademik disusun melalui mekanisme jelas", "3,62", "≥ 3,50", "Tercapai"],
        ["3", "Dosen dilibatkan dalam penyusunan/evaluasi kebijakan", "3,48", "≥ 3,50", "Belum"],
        ["4", "Pimpinan memberikan arahan yang jelas", "3,60", "≥ 3,50", "Tercapai"],
        ["5", "Pengambilan keputusan akademik objektif/transparan", "3,57", "≥ 3,50", "Tercapai"],
        ["6", "Informasi kebijakan akademik mudah diakses dosen", "3,65", "≥ 3,50", "Tercapai"],
        ["7", "Koordinasi pimpinan, fakultas, prodi berjalan baik", "3,50", "≥ 3,50", "Tercapai"],
        ["8", "Tata kelola mendukung pelaksanaan Tridharma", "3,48", "≥ 3,50", "Belum"],
        ["9", "Evaluasi pelaksanaan kebijakan akademik dilakukan", "3,60", "≥ 3,50", "Tercapai"],
        ["10", "Hasil evaluasi ditindaklanjuti untuk perbaikan", "3,52", "≥ 3,50", "Tercapai"],
        ["11", "Pimpinan bersikap terbuka terhadap masukan dosen", "3,49", "≥ 3,50", "Belum"],
        ["12", "Sistem tata kelola mendukung pencapaian VMTS", "3,71", "≥ 3,50", "Tercapai"],
        ["13", "Pembagian tugas dan kewenangan antarunit berjalan", "3,55", "≥ 3,50", "Tercapai"],
        ["14", "Tata kelola menerapkan prinsip transparansi", "3,62", "≥ 3,50", "Tercapai"],
        ["15", "Tata pamong universitas berjalan efektif", "3,66", "≥ 3,50", "Tercapai"],
    ],
    kekuatan_list=[
        "Tata pamong berjalan efektif dengan skor 3,66 (Sangat Baik)",
        "Sistem tata kelola mendukung pencapaian VMTS (skor 3,71)",
        "Informasi kebijakan akademik mudah diakses oleh dosen (skor 3,65)",
        "Pengambilan keputusan akademik objektif dan transparan (skor 3,57)",
        "Pimpinan memberikan arahan yang jelas (skor 3,60)",
        "Senat dan SPI aktif dalam pengawasan tata kelola",
        "40 MoU aktif dan 60 MoA aktif dengan berbagai mitra",
    ],
    kelemahan_list=[
        "Keterlibatan dosen dalam penyusunan/evaluasi kebijakan masih kurang (skor 3,48)",
        "Tata kelola belum sepenuhnya mendukung pelaksanaan Tridharma (skor 3,48)",
        "Pimpinan belum sepenuhnya terbuka terhadap masukan dosen (skor 3,49)",
        "Implementasi kerja sama perlu peningkatan (hanya 60% yang terimplementasi)",
        "Belum ada AMI khusus untuk Kerjasama",
    ],
    rekomendasi_list=[
        "Meningkatkan keterlibatan dosen dalam penyusunan dan evaluasi kebijakan akademik",
        "Memperkuat tata kelola untuk mendukung pelaksanaan Tridharma secara optimal",
        "Mendorong pimpinan untuk lebih terbuka terhadap masukan dari dosen",
        "Menyelenggarakan AMI khusus untuk Kerjasama guna evaluasi yang lebih komprehensif",
        "Meningkatkan implementasi kerja sama dari 60% menjadi 75% pada tahun 2026",
        "Mengembangkan sistem informasi manajemen kerja sama yang terintegrasi",
    ],
    ami_rata_skor="3,59/4 (Sangat Baik)",
    ami_status="12 Tercapai, 3 Belum Tercapai",
    ami_dokumen="AMI TPTK 2024-2025 (Kode: SPMI/PPM/MONEV/TPTK/2025) + AMI MBKM 2024-2025"
)

# ============================================================
# BAB V - KRITERIA 3: MAHASISWA
# ============================================================

generate_kriteria_bab(doc, 5, 3, "Mahasiswa",
    deskripsi_paragrafs=[
        "Kriteria 3 mengevaluasi pengelolaan mahasiswa di Universitas Tulungagung, mencakup penerimaan, pembinaan, pengembangan, dan layanan kemahasiswaan. Mahasiswa adalah pelanggan utama perguruan tinggi, sehingga pengelolaan mahasiswa yang baik menjadi prasyarat bagi tercapainya mutu pendidikan. Universitas Tulungagung memandang mahasiswa tidak hanya sebagai objek pendidikan, tetapi juga sebagai mitra dalam pengembangan akademik.",
        "Penerimaan mahasiswa di Universitas Tulungagung dilakukan melalui mekanisme yang transparan dan akuntabel, dengan jalur SNMPTN, SBMPTN, dan mandiri. Universitas Tulungagung memastikan akses pendidikan bagi mahasiswa dari berbagai latar belakang ekonomi, melalui program beasiswa yang beragam. Selain itu, Universitas Tulungagung juga menerima mahasiswa transfer dan mahasiswa internasional.",
        "Universitas Tulungagung mengembangkan organisasi kemahasiswaan yang aktif dan demokratis, sebagai sarana pengembangan kepemimpinan dan soft skills mahasiswa. Organisasi kemahasiswaan mencakup BEM, HIMA, UKM, dan organisasi lainnya, dengan pembinaan dari Wakil Rektor III dan staf. Universitas Tulungagung juga menyediakan layanan karir untuk mendukung kesiapan kerja lulusan, melalui pusat karir yang aktif.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI Mahasiswa 2024/2025, seluruh 14 indikator berstatus Tercapai. Angka Efisiensi Edukasi (AEE PT) mencapai 92,50% (target ≥90%), persentase lulusan bekerja/berwirausaha/melanjutkan studi mencapai 84,30% (target ≥80%), dan persentase mahasiswa berkegiatan organisasi mencapai 68,40% (target ≥60%).",
        "Program penguatan integritas akademik dan adaptasi mencapai 95,00%, pengukuran capaian standar kompetensi lulusan mencapai 91,80%, dan kontribusi mahasiswa dalam pencapaian SDGs mencapai 76,20%. Pengakuan dan apresiasi kompetensi lulusan oleh DUDI mencapai 82,40%, sementara perbaikan proses pembelajaran berdasarkan hasil AMI mencapai 94,50%.",
        "Hasil ini menunjukkan bahwa pengelolaan mahasiswa di Universitas Tulungagung berjalan dengan baik, dengan semua indikator tercapai. Namun, beberapa indikator seperti kontribusi mahasiswa dalam SDGs (76,20%) dan pengakuan DUDI (82,40%) masih dapat ditingkatkan untuk mencapai target yang lebih tinggi.",
    ],
    ami_data_table=[
        ["1", "Angka Efisiensi Edukasi (AEE PT)", "92,50%", "≥ 90%", "Tercapai"],
        ["2", "Lulusan bekerja/berwirausaha/melanjutkan studi", "84,30%", "≥ 80%", "Tercapai"],
        ["3", "Mahasiswa berkegiatan/berorganisasi", "68,40%", "≥ 60%", "Tercapai"],
        ["4", "Kebijakan dan pedoman penerimaan mahasiswa", "100%", "100%", "Tercapai"],
        ["5", "Sistem pengelolaan dan layanan kemahasiswaan", "100%", "100%", "Tercapai"],
        ["6", "Layanan pembelajaran di luar program reguler", "100%", "100%", "Tercapai"],
        ["7", "Program penguatan integritas akademik", "95,00%", "≥ 90%", "Tercapai"],
        ["8", "Analisis prestasi mahasiswa, keterserapan lulusan", "100%", "100%", "Tercapai"],
        ["9", "Pengukuran capaian standar kompetensi lulusan", "91,80%", "≥ 90%", "Tercapai"],
        ["10", "Implementasi Pendidikan Anti Korupsi", "100%", "100%", "Tercapai"],
        ["11", "Program pencegahan dan penanganan kekerasan", "100%", "100%", "Tercapai"],
        ["12", "Pengakuan dan apresiasi kompetensi lulusan DUDI", "82,40%", "≥ 80%", "Tercapai"],
        ["13", "Perbaikan proses pembelajaran berdasarkan AMI", "94,50%", "≥ 90%", "Tercapai"],
        ["14", "Kontribusi mahasiswa dalam pencapaian SDGs", "76,20%", "≥ 70%", "Tercapai"],
    ],
    kekuatan_list=[
        "Seluruh 14 indikator AMI Mahasiswa berstatus Tercapai",
        "AEE PT mencapai 92,50%, di atas target 90%",
        "84,30% lulusan bekerja/berwirausaha/melanjutkan studi",
        "Program penguatan integritas akademik mencapai 95%",
        "Implementasi Pendidikan Anti Korupsi 100%",
        "Sistem pengelolaan dan layanan kemahasiswaan tersedia 100%",
    ],
    kelemahan_list=[
        "Kontribusi mahasiswa dalam SDGs masih 76,20%, perlu peningkatan",
        "Pengakuan kompetensi lulusan oleh DUDI masih 82,40%, perlu peningkatan",
        "Mahasiswa internasional masih terbatas",
        "Tracer study response rate masih perlu ditingkatkan",
    ],
    rekomendasi_list=[
        "Meningkatkan kontribusi mahasiswa dalam SDGs melalui program PkM terstruktur",
        "Memperkuat hubungan dengan DUDI untuk peningkatan apresiasi kompetensi lulusan",
        "Mengembangkan program mahasiswa internasional melalui kerja sama PT luar negeri",
        "Meningkatkan response rate tracer study melalui insentif dan follow-up",
        "Mengembangkan pusat karir yang lebih aktif dengan job fair berkala",
        "Menyediakan beasiswa untuk mahasiswa berprestasi dan kelompok rentan",
    ],
    ami_rata_skor="14/14 Tercapai (100%)",
    ami_status="Semua indikator Tercapai",
    ami_dokumen="AMI Mahasiswa 2024-2025 (Kode: SPMI/PPM/MONEV/MHS/2025)"
)

# ============================================================
# BAB VI - KRITERIA 4: SDM
# ============================================================

generate_kriteria_bab(doc, 6, 4, "Sumber Daya Manusia",
    deskripsi_paragrafs=[
        "Kriteria 4 mengevaluasi pengelolaan Sumber Daya Manusia (SDM) di Universitas Tulungagung, mencakup dosen dan tenaga kependidikan. SDM yang kompeten menjadi prasyarat bagi penyelenggaraan Tridharma yang bermutu. Pengelolaan SDM mencakup perencanaan, rekrutmen, pengembangan, penilaian, dan retensi. Universitas Tulungagung mengembangkan sistem manajemen SDM yang berbasis kinerja.",
        "Berdasarkan data terbaru Tahun 2025, Universitas Tulungagung memiliki total 93 orang SDM yang terdiri dari dosen dan tenaga kependidikan. Kualifikasi akademik minimal S2 bagi dosen telah mencapai 100%, sementara kualifikasi minimal D3 untuk tenaga kependidikan juga telah mencapai 100%. Sebanyak 78,40% dosen tetap memiliki jabatan fungsional minimal Lektor.",
        "Kepemilikan sertifikat pendidik telah mencapai 94,60% dari target 100%, sementara persentase dosen tidak tetap dengan sertifikat kompetensi berada pada kisaran 16,20% dari total dosen, yang masih di bawah batas maksimum 20%. Realisasi anggaran pengembangan SDM mencapai 88,70% dari total kebutuhan, berada dalam rentang target 85-95%.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI SDM 2024/2025, dari 18 indikator yang dievaluasi, 16 indikator berstatus Tercapai dan 2 indikator berstatus Belum Tercapai. Dua indikator yang belum tercapai adalah: persentase dosen tetap dengan jabatan fungsional minimal Lektor (52,40% dari target ≥75%) dan kepemilikan sertifikat pendidik oleh dosen (94,60% dari target 100%).",
        "Indikator yang tercapai mencakup: kualifikasi akademik minimal S2 bagi dosen (100%), kualifikasi minimal D3 untuk tendik (100%), keanggotaan tenaga perpustakaan profesional (100%), pelaksanaan uji kompetensi (100%), ketersediaan laporan monev kinerja dosen dan tendik (100%), ketersediaan dokumen DP3 (100%), ketersediaan dokumen Renstra/Renop SDM (100%), terpenuhinya biaya pengembangan SDM (88,70%), ketersediaan Pedoman Kesejahteraan (100%), pelaksanaan survei kepuasan (100%), efektivitas sistem rekrutmen (95,20%), dan ketersediaan struktur organisasi pengelolaan SDM (100%).",
        "Selain AMI SDM, hasil AMI Kesejahteraan 2024/2025 menunjukkan rata-rata capaian 96,32% (Sangat Baik) dari 25 indikator, yang seluruhnya berstatus Tercapai. Ini menunjukkan bahwa kesejahteraan SDM di Universitas Tulungagung berjalan dengan sangat baik, mencakup buku pedoman, SOP, hak karyawan, imbal jasa, asuransi, cuti, dana pensiun, dan program rekreasi.",
    ],
    ami_data_table=[
        ["1", "Kualifikasi akademik minimal S2 bagi dosen", "100%", "100%", "Tercapai"],
        ["2", "Dosen tetap dengan jabatan fungsional minimal Lektor", "52,40%", "≥ 75%", "Belum"],
        ["3", "Kepemilikan sertifikat pendidik dosen", "94,60%", "100%", "Belum"],
        ["4", "Dosen tidak tetap dengan sertifikat kompetensi", "16,20%", "< 20%", "Tercapai"],
        ["5", "Kualifikasi akademik minimal D3 tendik", "100%", "100%", "Tercapai"],
        ["6", "Keanggotaan tenaga perpustakaan profesional", "100%", "100%", "Tercapai"],
        ["7", "Pelaksanaan uji kompetensi keahlian", "100%", "100%", "Tercapai"],
        ["8", "Ketersediaan laporan monev kinerja dosen", "100%", "100%", "Tercapai"],
        ["9", "Ketersediaan laporan monev kinerja tendik", "100%", "100%", "Tercapai"],
        ["10", "Ketersediaan dokumen Laporan DP3", "100%", "100%", "Tercapai"],
        ["11", "Ketersediaan dokumen Renstra/Renop SDM", "100%", "100%", "Tercapai"],
        ["12", "Terpenuhinya biaya pengembangan SDM", "88,70%", "85-95%", "Tercapai"],
        ["13", "Ketersediaan Pedoman Kesejahteraan", "100%", "100%", "Tercapai"],
        ["14", "Pelaksanaan survei kepuasan kesejahteraan", "100%", "100%", "Tercapai"],
        ["15", "Penyempurnaan pedoman kesejahteraan", "100%", "100%", "Tercapai"],
        ["16", "Efektivitas sistem rekrutmen SDM", "95,20%", "≥ 90%", "Tercapai"],
        ["17", "Ketersediaan lembaga pengembangan SDM", "100%", "100%", "Tercapai"],
        ["18", "Ketersediaan struktur organisasi SDM", "100%", "100%", "Tercapai"],
    ],
    kekuatan_list=[
        "16 dari 18 indikator AMI SDM berstatus Tercapai",
        "Kualifikasi akademik dosen (S2) dan tendik (D3) mencapai 100%",
        "AMI Kesejahteraan mencapai rata-rata 96,32% (Sangat Baik), semua tercapai",
        "Sistem rekrutmen efektif (95,20%)",
        "Ketersediaan dokumen monev, DP3, Renstra/Renop SDM mencapai 100%",
        "Realisasi anggaran pengembangan SDM 88,70% (dalam target 85-95%)",
    ],
    kelemahan_list=[
        "Jabatan fungsional Lektor+ baru 52,40% (target ≥75%)",
        "Sertifikat pendidik dosen baru 94,60% (target 100%)",
        "Total SDM masih 93 orang, perlu peningkatan seiring pertumbuhan mahasiswa",
        "Program studi lanjut (S3) bagi dosen masih terbatas",
    ],
    rekomendasi_list=[
        "Mendorong dosen untuk promosi jabatan akademik ke Lektor dan Lektor Kepala",
        "Memfasilitasi dosen yang belum bersertifikat untuk mengikuti sertifikasi pendidik",
        "Menambah jumlah SDM melalui rekrutmen dosen dan tendik baru",
        "Menyediakan beasiswa studi lanjut S3 bagi dosen",
        "Meningkatkan program pelatihan kompetensi bagi dosen dan tendik",
        "Mengembangkan sistem informasi SDM yang terintegrasi dengan PD Dikti",
    ],
    ami_rata_skor="16/18 Tercapai (88,9%) + Kesejahteraan 96,32%",
    ami_status="16 Tercapai, 2 Belum Tercapai",
    ami_dokumen="AMI SDM 2024-2025 + AMI Kesejahteraan 2024-2025"
)

# ============================================================
# BAB VII - KRITERIA 5: KEUANGAN, SARANA, PRASARANA
# ============================================================

generate_kriteria_bab(doc, 7, 5, "Keuangan, Sarana, dan Prasarana",
    deskripsi_paragrafs=[
        "Kriteria 5 mengevaluasi pengelolaan keuangan, sarana, dan prasarana di Universitas Tulungagung secara terintegrasi. Pengelolaan yang baik dari ketiga aspek ini menjadi prasyarat bagi keberlanjutan institusi. Universitas Tulungagung menerapkan prinsip-prinsip tata kelola yang baik dalam pengelolaan keuangan, sarana, dan prasarana, dengan audit internal dan eksternal secara berkala.",
        "Pada Tahun Anggaran 2025, total pendapatan Universitas Tulungagung yang ditargetkan sesuai RENOP mencapai Rp28.400 juta, dengan komposisi pendapatan dari pemerintah sebesar Rp3.570 juta, pendapatan dari mahasiswa sebesar Rp10.405 juta, pendapatan kegiatan profesional sebesar Rp6.977 juta, dan pendapatan sumber lainnya sebesar Rp7.446 juta. Total belanja operasional pendidikan diproyeksikan sebesar Rp26.390 juta.",
        "Untuk sarana dan prasarana, berdasarkan hasil AMI 2024/2025, persentase sarana dalam kondisi baik/terawat mencapai 95,20% (dari 83% di tahun 2023), sementara prasarana mencapai 95,80% (dari 84%). Peningkatan signifikan ini menunjukkan komitmen Universitas Tulungagung dalam investasi sarana dan prasarana.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI Keuangan 2024/2025, seluruh 14 indikator berstatus Tercapai. Realisasi pendapatan terhadap target RENOP mencapai 95,20%, realisasi belanja operasional pendidikan 94,80%, belanja kemahasiswaan 95,60%, belanja penelitian 96,10%, dan belanja pengabdian 94,30%. Rasio pendapatan terhadap belanja mencapai 97,30% (target ≥95%).",
        "Untuk sarana dan prasarana, seluruh 14 indikator AMI berstatus Tercapai. Persentase sarana baik mencapai 95,20%, prasarana baik 95,80%, ruang kelas sesuai standar 94,50%, laboratorium sesuai standar 96,30%, ruang dosen sesuai standar 95,70%, perpustakaan 97,20%, fasilitas ICT 95,40%, dan fasilitas kemahasiswaan 94,80%.",
        "Hasil ini menunjukkan bahwa pengelolaan keuangan, sarana, dan prasarana di Universitas Tulungagung berjalan dengan sangat baik, dengan semua indikator tercapai. Peningkatan signifikan terjadi pada sarana dan prasarana dibandingkan tahun 2023, menunjukkan efektivitas investasi yang dilakukan.",
    ],
    ami_data_table=[
        ["1", "Realisasi Pendapatan terhadap Target RENOP", "95,20%", "≥ 90%", "Tercapai"],
        ["2", "Realisasi Belanja Operasional Pendidikan", "94,80%", "≥ 90%", "Tercapai"],
        ["3", "Realisasi Belanja Kemahasiswaan", "95,60%", "≥ 90%", "Tercapai"],
        ["4", "Realisasi Belanja Penelitian", "96,10%", "≥ 90%", "Tercapai"],
        ["5", "Realisasi Belanja Pengabdian", "94,30%", "≥ 90%", "Tercapai"],
        ["6", "Realisasi Belanja Investasi SDM", "95,70%", "≥ 90%", "Tercapai"],
        ["7", "Realisasi Belanja Investasi Sarana", "95,40%", "≥ 90%", "Tercapai"],
        ["8", "Realisasi Belanja Investasi Prasarana", "94,90%", "≥ 90%", "Tercapai"],
        ["9", "Rasio Pendapatan terhadap Belanja", "97,30%", "≥ 95%", "Tercapai"],
        ["10", "Sarana dalam kondisi baik/terawat", "95,20%", "≥ 90%", "Tercapai"],
        ["11", "Prasarana dalam kondisi baik/terawat", "95,80%", "≥ 90%", "Tercapai"],
        ["12", "Ruang kelas sesuai standar rasio", "94,50%", "≥ 90%", "Tercapai"],
        ["13", "Laboratorium sesuai standar", "96,30%", "≥ 90%", "Tercapai"],
        ["14", "Fasilitas ICT, internet, WiFi", "95,40%", "≥ 90%", "Tercapai"],
    ],
    kekuatan_list=[
        "Seluruh 14 indikator AMI Keuangan berstatus Tercapai",
        "Seluruh 14 indikator AMI Sarana Prasarana berstatus Tercapai",
        "Realisasi pendapatan 95,20% dari target RENOP",
        "Sarana baik meningkat dari 83% (2023) menjadi 95,20% (2025)",
        "Prasarana baik meningkat dari 84% (2023) menjadi 95,80% (2025)",
        "Rasio pendapatan terhadap belanja mencapai 97,30%",
        "Total pendapatan target Rp28.400 juta",
    ],
    kelemahan_list=[
        "Pendapatan dari sumber selain SPP masih perlu diversifikasi",
        "Beberapa sarana (scanner, stand LCD, diesel) perlu perbaikan",
        "Opini audit eksternal masih WDP, target WTP",
        "Smart campus components masih terbatas",
    ],
    rekomendasi_list=[
        "Meningkatkan diversifikasi sumber pendapatan melalui RGU dan kerja sama DUDI",
        "Memperbaiki sarana yang dalam kondisi kurang baik (scanner, LCD, diesel)",
        "Meningkatkan tata kelola keuangan untuk mencapai opini WTP",
        "Mengembangkan smart campus dengan WiFi campus-wide dan smart classroom",
        "Menginvestasikan dana untuk pemeliharaan sarana-prasarana secara berkala",
        "Mengembangkan sistem inventarisasi berbasis IT untuk monitoring real-time",
    ],
    ami_rata_skor="14/14 Keuangan + 14/14 Sarana (100% Tercapai)",
    ami_status="Semua indikator Tercapai",
    ami_dokumen="AMI Keuangan 2024-2025 + AMI Sarana Prasarana 2024-2025"
)

# ============================================================
# BAB VIII - KRITERIA 6: PENDIDIKAN
# ============================================================

generate_kriteria_bab(doc, 8, 6, "Pendidikan",
    deskripsi_paragrafs=[
        "Kriteria 6 mengevaluasi penyelenggaraan pendidikan di Universitas Tulungagung, mencakup Standar Kompetensi Lulusan, Standar Proses Pembelajaran, Standar Penilaian, Standar Pengelolaan, Standar Isi, Standar Dosen dan Tenaga Kependidikan, Standar Sarana dan Prasarana, serta Standar Pembiayaan. Kriteria ini merupakan kriteria dengan evaluasi paling komprehensif, mencakup 19 indikator dari AMI Pendidikan dan 20 indikator dari AMI Proses Pembelajaran.",
        "Berdasarkan hasil AMI Pendidikan 2024/2025, terdapat peningkatan signifikan dibandingkan periode AMI 2023. Beberapa capaian yang telah ditingkatkan meliputi: tingkat penyerapan kerja lulusan dari 80% menjadi 86,50%, peningkatan persentase lulusan wirausahawan dari 15% menjadi 18,20%, perbaikan rata-rata IPK dari 3,42 menjadi 3,48, percepatan masa studi, dan peningkatan rasio dosen pembimbing tugas akhir dari 12:1 menjadi 8:1.",
        "Untuk proses pembelajaran, hasil AMI menunjukkan tren peningkatan skor rata-rata dari 3,43 (Gasal 2023-2) menjadi 3,49 (Gasal 2024-1) menjadi 3,59 (Genap 2024-2), dengan seluruh fakultas (FE, FH, FISIP, FP, FT, F.KES) menunjukkan peningkatan performa yang substansial. Peningkatan ini merupakan hasil konsisten dari implementasi Rencana Tindak Lanjut (RTL) periode sebelumnya.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI Pendidikan 2024/2025, seluruh 19 indikator berstatus Tercapai. Capaian utama meliputi: persentase lulusan terserap di dunia kerja 86,50% (target ≥70%), persentase lulusan wirausahawan 18,20% (target ≥10%), rata-rata IPK lulusan 3,48 (target ≥3,00), rata-rata masa studi 3,8 tahun (target ≤4,5 tahun), dan rata-rata masa tunggu kerja 3,2 bulan (target ≤6 bulan).",
        "Untuk standar isi, ketersediaan dan kelengkapan dokumen kurikulum mencapai 100%, frekuensi tinjauan kurikulum 2 kali/tahun, rasio mahasiswa per dosen pembimbing akademik 16:1 (target ≤20:1), dan rasio mahasiswa per dosen pembimbing tugas akhir 8:1 (target ≤10:1). Persentase mata kuliah blended learning mencapai 32% (target ≥20%), dan rata-rata jumlah mahasiswa per kelas 32 (target ≤40).",
        "Untuk penilaian, persentase dosen menerapkan berbagai metode penilaian mencapai 100%, ketepatan waktu input nilai pasca ujian mencapai 100%, dan persentase dosen berkualifikasi S2/S3 mencapai 100%. Persentase dosen dan tendik bersertifikat pendidik mencapai 88,50% (target ≥75%). Ketersediaan fasilitas pembelajaran sesuai standar mencapai 100%.",
        "Untuk proses pembelajaran, dari 20 indikator AMI Proses Pembelajaran (Genap 2024-2), seluruhnya berstatus Tercapai dengan skor di atas 3,50. Skor tertinggi adalah penguasaan materi perkuliahan oleh dosen (3,82), diikuti oleh penyampaian RPS pada pertemuan pertama (3,78), tahapan perkuliahan (3,77), dan metode pembelajaran efektif (3,76).",
    ],
    ami_data_table=[
        ["1", "Lulusan terserap di dunia kerja", "86,50%", "≥ 70%", "Tercapai"],
        ["2", "Lulusan menjadi wirausahawan", "18,20%", "≥ 10%", "Tercapai"],
        ["3", "Rata-rata IPK lulusan", "3,48", "≥ 3,00", "Tercapai"],
        ["4", "Rata-rata masa studi lulusan", "3,8 thn", "≤ 4,5 thn", "Tercapai"],
        ["5", "Rata-rata masa tunggu kerja", "3,2 bln", "≤ 6 bln", "Tercapai"],
        ["6", "Dokumen kurikulum tersedia lengkap", "100%", "100%", "Tercapai"],
        ["7", "Frekuensi tinjauan kurikulum", "2x/thn", "≥ 1x/thn", "Tercapai"],
        ["8", "Rasio mhs/dosen pembimbing akademik", "16:1", "≤ 20:1", "Tercapai"],
        ["9", "Rasio mhs/dosen pembimbing TA", "8:1", "≤ 10:1", "Tercapai"],
        ["10", "Bahan ajar dan referensi tersedia", "100%", "100%", "Tercapai"],
        ["11", "Mata kuliah blended learning", "32%", "≥ 20%", "Tercapai"],
        ["12", "Rata-rata mahasiswa per kelas", "32", "≤ 40", "Tercapai"],
        ["13", "Monitoring kehadiran dan pembelajaran", "2x/smt", "≥ 1x/smt", "Tercapai"],
        ["14", "Dosen menerapkan berbagai metode penilaian", "100%", "100%", "Tercapai"],
        ["15", "Ketepatan waktu input nilai", "100%", "100%", "Tercapai"],
        ["16", "Dosen berkualifikasi S2/S3", "100%", "100%", "Tercapai"],
        ["17", "Dosen/tendik bersertifikat pendidik", "88,50%", "≥ 75%", "Tercapai"],
        ["18", "Fasilitas pembelajaran sesuai standar", "100%", "100%", "Tercapai"],
        ["19", "Dokumen RPS dan laporan pelaksanaan", "100%", "100%", "Tercapai"],
    ],
    kekuatan_list=[
        "Seluruh 19 indikator AMI Pendidikan berstatus Tercapai",
        "Seluruh 20 indikator AMI Proses Pembelajaran berstatus Tercapai (skor 3,59)",
        "Lulusan terserap DUDI 86,50% (target ≥70%)",
        "IPK rata-rata lulusan 3,48 (target ≥3,00)",
        "Masa studi 3,8 tahun (target ≤4,5 tahun), masa tunggu 3,2 bulan (target ≤6 bulan)",
        "Dokumen kurikulum 100% lengkap, tinjauan 2x/tahun",
        "Rasio dosen pembimbing TA 8:1 (dari 12:1, target ≤10:1)",
        "Skor proses pembelajaran Genap 3,59 (dari Ganjil 3,49, dari 2023-2: 3,43)",
    ],
    kelemahan_list=[
        "Persentase dosen bersertifikat pendidik masih 88,50% (target 100%)",
        "Blended learning baru 32% (target ≥20%, namun masih bisa ditingkatkan)",
        "Metode pembelajaran daring masih skor terendah (3,65), meski di atas target",
        "Integrasi hasil penelitian ke pembelajaran masih bisa ditingkatkan",
    ],
    rekomendasi_list=[
        "Memfasilitasi dosen yang belum bersertifikat untuk mengikuti sertifikasi pendidik",
        "Meningkatkan persentase mata kuliah blended learning dari 32% ke 40% pada 2026",
        "Mengembangkan pelatihan metode pembelajaran daring yang lebih efektif",
        "Mengintegrasikan hasil penelitian dan PkM ke dalam materi pembelajaran",
        "Mengembangkan LMS yang lebih komprehensif dengan fitur interaktif",
        "Menyelenggarakan workshop penyusunan RPS berbasis OBE secara berkala",
        "Mengembangkan peer teaching observation antar-dosen untuk sharing best practice",
    ],
    ami_rata_skor="19/19 Pendidikan + 20/20 Proses Pembelajaran (100% Tercapai)",
    ami_status="Semua indikator Tercapai",
    ami_dokumen="AMI Pendidikan 2024-2025 + AMI Proses Pembelajaran Ganjil & Genap 2024-2025"
)

# ============================================================
# BAB IX - KRITERIA 7: PENELITIAN
# ============================================================

generate_kriteria_bab(doc, 9, 7, "Penelitian",
    deskripsi_paragrafs=[
        "Kriteria 7 mengevaluasi penyelenggaraan penelitian di Universitas Tulungagung, mencakup standar hasil penelitian, standar isi, standar proses, standar penilaian, standar pelaksana, standar sarana, standar pengelolaan, dan standar pembiayaan penelitian. Penelitian menjadi salah satu pilar Tridharma Perguruan Tinggi yang berkontribusi pada pengembangan IPTEK dan kesejahteraan masyarakat.",
        "Universitas Tulungagung melalui Lembaga Penelitian dan Pengabdian kepada Masyarakat (LPPM) mengelola pelaksanaan penelitian dengan mengacu pada standar penelitian sebagaimana diatur dalam Permen 39 Tahun 2025. LPPM bertanggung jawab atas perencanaan, pelaksanaan, monitoring, evaluasi, dan pelaporan penelitian, dengan dukungan UPPS dan GKM di tingkat program studi.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI Penelitian 2024/2025, seluruh 20 indikator berstatus Tercapai. Ketersediaan panduan penelitian mencapai 100%, sosialisasi panduan 100%, laporan akhir penelitian 96,80%, integrasi hasil penelitian ke RPS 93,50%, kesesuaian dengan roadmap 97,20%, publikasi ilmiah 95,60%, HKI 93,40%, sitasi 90,80%, dan adopsi industri 88,50%.",
        "Proporsi skema penelitian (Dasar dan Terapan) mencapai 96,80%, jumlah penelitian sesuai roadmap 93,40%, proposal sesuai roadmap 95,20%, laporan kemajuan 97,80%, laporan akhir 96,50%, panduan penilaian 100%, laporan penilaian 100%, laporan Monev 100%, laporan RTL 93,80%, dan jumlah dosen peneliti 91,76%.",
        "Hasil ini menunjukkan bahwa penelitian di Universitas Tulungagung berjalan dengan baik, dengan semua indikator tercapai. Namun, beberapa indikator seperti adopsi industri (88,50%) dan jumlah dosen peneliti (91,76%) masih dapat ditingkatkan untuk mencapai capaian yang lebih optimal.",
    ],
    ami_data_table=[
        ["1", "Ketersediaan panduan penelitian", "100%", "Tersedia", "Tercapai"],
        ["2", "Sosialisasi panduan penelitian", "100%", "Terlaksana", "Tercapai"],
        ["3", "Laporan akhir penelitian tersedia", "96,80%", "90-100%", "Tercapai"],
        ["4", "Integrasi hasil penelitian ke RPS", "93,50%", "≥10-15%", "Tercapai"],
        ["5", "Kesesuaian dengan roadmap visi-misi", "97,20%", "75-85%", "Tercapai"],
        ["6", "Publikasi ilmiah dosen per prodi", "95,60%", "10 dok/prodi", "Tercapai"],
        ["7", "Perolehan HKI (Paten, Hak Cipta)", "93,40%", "Meningkat", "Tercapai"],
        ["8", "Sitasi karya ilmiah dan buku", "90,80%", "Meningkat", "Tercapai"],
        ["9", "Adopsi industri/hasil penelitian", "88,50%", "Ada/thn", "Tercapai"],
        ["10", "Sosialisasi dan pelatihan penelitian", "100%", "Terlaksana", "Tercapai"],
        ["11", "Proporsi skema penelitian", "96,80%", "Sesuai", "Tercapai"],
        ["12", "Penelitian sesuai roadmap", "93,40%", "Progresif", "Tercapai"],
        ["13", "Proposal sesuai roadmap", "95,20%", "Sesuai", "Tercapai"],
        ["14", "Laporan kemajuan penelitian", "97,80%", "100%", "Tercapai"],
        ["15", "Keterpenuhan laporan akhir", "96,50%", "90-100%", "Tercapai"],
        ["16", "Panduan penilaian proposal", "100%", "Tersedia", "Tercapai"],
        ["17", "Laporan penilaian pelaksanaan", "100%", "Tersedia", "Tercapai"],
        ["18", "Laporan Monitoring dan Evaluasi", "100%", "Tersedia", "Tercapai"],
        ["19", "Laporan Rencana Tindak Lanjut", "93,80%", "Tersedia", "Tercapai"],
        ["20", "Dosen yang memenuhi kriteria peneliti", "91,76%", "Meningkat", "Tercapai"],
    ],
    kekuatan_list=[
        "Seluruh 20 indikator AMI Penelitian berstatus Tercapai",
        "Panduan penelitian dan sosialisasi mencapai 100%",
        "Kesesuaian dengan roadmap 97,20%",
        "Laporan kemajuan 97,80%, laporan akhir 96,50%",
        "Publikasi ilmiah 95,60% (10 dokumen/prodi/tahun)",
        "Panduan penilaian, laporan Monev, laporan RTL tersedia 100%",
    ],
    kelemahan_list=[
        "Adopsi industri masih 88,50%, perlu peningkatan",
        "Jumlah dosen peneliti 91,76%, perlu peningkatan",
        "Sitasi karya ilmiah 90,80%, masih bisa ditingkatkan",
        "Integrasi hasil penelitian ke RPS 93,50%, masih bisa ditingkatkan",
    ],
    rekomendasi_list=[
        "Meningkatkan adopsi industri melalui kerja sama DUDI untuk penelitian terapan",
        "Mendorong lebih banyak dosen untuk menjadi peneliti aktif",
        "Meningkatkan jumlah publikasi di jurnal internasional bereputasi (Scopus)",
        "Mengintegrasikan hasil penelitian ke dalam materi pembelajaran (RBL)",
        "Mengembangkan kerja sama penelitian dengan PT lain dan lembaga penelitian",
        "Menyediakan insentif bagi dosen yang menghasilkan publikasi dan HKI",
    ],
    ami_rata_skor="20/20 Tercapai (100%)",
    ami_status="Semua indikator Tercapai",
    ami_dokumen="AMI Penelitian 2024-2025 (Kode: SPMI/PPM/MONEV/PEN/2025)"
)

# ============================================================
# BAB X - KRITERIA 8: PkM
# ============================================================

generate_kriteria_bab(doc, 10, 8, "Pengabdian kepada Masyarakat",
    deskripsi_paragrafs=[
        "Kriteria 8 mengevaluasi penyelenggaraan Pengabdian kepada Masyarakat (PkM) di Universitas Tulungagung. PkM merupakan wujud nyata kontribusi perguruan tinggi kepada masyarakat, sebagaimana diamanatkan dalam Tujuan Universitas Tulungagung: terwujudnya pengabdian kepada masyarakat berbasis penalaran dan karya penelitian yang bermanfaat dalam memajukan kesejahteraan umum dan mencerdaskan kehidupan bangsa.",
        "Universitas Tulungagung melalui LPPM mengelola pelaksanaan PkM dengan mengacu pada standar PkM sebagaimana diatur dalam Permen 39 Tahun 2025. PkM mencakup 8 standar: standar hasil, standar isi, standar proses, standar penilaian, standar pelaksana, standar sarana, standar pengelolaan, dan standar pembiayaan PkM.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI Pengabdian kepada Masyarakat 2024/2025, seluruh 20 indikator berstatus Tercapai. Ketersediaan panduan PkM mencapai 100%, sosialisasi 100%, laporan akhir 96,20%, integrasi ke RPS 92,80%, kesesuaian dengan roadmap 95,40%, publikasi PkM 94,50%, HKI/TTG 93,80%, penyelesaian masalah mitra 96,80%, dan adopsi hasil PkM 91,20%.",
        "Hasil AMI periode sebelumnya (2023/2024) menunjukkan bahwa 12 dari 13 indikator PkM berada dalam kategori terpenuhi, namun capaian Teknologi Tepat Guna (TTG) masih rendah (20%). Pada AMI 2024/2025, TTG telah meningkat menjadi 93,80%, menunjukkan perbaikan signifikan dalam pengembangan teknologi tepat guna.",
    ],
    ami_data_table=[
        ["1", "Ketersediaan panduan PkM", "100%", "Tersedia", "Tercapai"],
        ["2", "Sosialisasi panduan PkM", "100%", "Terlaksana", "Tercapai"],
        ["3", "Laporan akhir PkM tersedia", "96,20%", "90-100%", "Tercapai"],
        ["4", "Integrasi hasil PkM ke RPS", "92,80%", "≥10-15%", "Tercapai"],
        ["5", "Kesesuaian PkM dengan roadmap", "95,40%", "75-85%", "Tercapai"],
        ["6", "Publikasi ilmiah dari hasil PkM", "94,50%", "Meningkat", "Tercapai"],
        ["7", "HKI, Teknologi Tepat Guna (TTG)", "93,80%", "Meningkat", "Tercapai"],
        ["8", "Penyelesaian masolah mitra/desa", "96,80%", "Aktif", "Tercapai"],
        ["9", "Adopsi hasil PkM oleh masyarakat", "91,20%", "Ada/thn", "Tercapai"],
        ["10", "Sosialisasi dan pelatihan PkM", "100%", "Terlaksana", "Tercapai"],
        ["11", "Proporsi skema PkM", "95,60%", "Sesuai", "Tercapai"],
        ["12", "Kegiatan PkM sesuai roadmap", "94,20%", "Progresif", "Tercapai"],
        ["13", "Proposal PkM sesuai roadmap", "95,80%", "Sesuai", "Tercapai"],
        ["14", "Laporan kemajuan kegiatan PkM", "97,60%", "100%", "Tercapai"],
        ["15", "Keterpenuhan laporan akhir PkM", "96,20%", "90-100%", "Tercapai"],
        ["16", "Panduan penilaian proposal PkM", "100%", "Tersedia", "Tercapai"],
        ["17", "Laporan penilaian pelaksanaan PkM", "100%", "Tersedia", "Tercapai"],
        ["18", "Laporan Monitoring dan Evaluasi PkM", "100%", "Tersedia", "Tercapai"],
        ["19", "Laporan Rencana Tindak Lanjut PkM", "93,50%", "Tersedia", "Tercapai"],
        ["20", "Dosen yang memenuhi kriteria pengabdi", "92,40%", "Meningkat", "Tercapai"],
    ],
    kekuatan_list=[
        "Seluruh 20 indikator AMI PkM berstatus Tercapai",
        "TTG meningkat dari 20% (2023) menjadi 93,80% (2025)",
        "Penyelesaian masalah mitra 96,80%",
        "Laporan kemajuan PkM 97,60%",
        "Panduan PkM dan sosialisasi 100%",
        "Adopsi hasil PkM oleh masyarakat 91,20%",
    ],
    kelemahan_list=[
        "Integrasi hasil PkM ke RPS masih 92,80%, perlu peningkatan",
        "Adopsi hasil PkM masih 91,20%, perlu peningkatan",
        "Dosen pengabdi 92,40%, perlu peningkatan",
        "Diseminasi PkM melalui forum nasional/internasional masih terbatas",
    ],
    rekomendasi_list=[
        "Meningkatkan integrasi hasil PkM ke dalam materi pembelajaran (CBL)",
        "Mengembangkan produk PkM yang dapat dikomersialisasi",
        "Mendorong lebih banyak dosen untuk aktif dalam PkM",
        "Menyelenggarakan seminar nasional PkM tahunan",
        "Mengembangkan kerja sama PkM dengan pemerintah daerah Jawa Timur",
        "Meningkatkan publikasi PkM di jurnal terakreditasi SINTA",
    ],
    ami_rata_skor="20/20 Tercapai (100%)",
    ami_status="Semua indikator Tercapai",
    ami_dokumen="AMI Pengabdian kepada Masyarakat 2024-2025 (Kode: SPMI/PPM/MONEV/PKM/2025)"
)

# ============================================================
# BAB XI - KRITERIA 9: LUARAN DAN CAPAIAN TRIDHARMA
# ============================================================

generate_kriteria_bab(doc, 11, 9, "Luaran dan Capaian Tridharma",
    deskripsi_paragrafs=[
        "Kriteria 9 mengevaluasi luaran dan capaian Tridharma Universitas Tulungagung secara terintegrasi. Luaran Tridharma mencakup lulusan, publikasi, paten, produk, jasa, dan kontribusi nyata kepada masyarakat. Universitas Tulungagung menargetkan peningkatan kuantitas dan kualitas luaran Tridharma, sebagaimana tertuang dalam sasaran strategis dan Indikator Kinerja Utama (IKU).",
        "Luaran Tridharma diukur melalui 8 Indikator Kinerja Utama (IKU) yang diturunkan dari Renstra. IKU mencakup: (1) lulusan terserap DUDI; (2) mahasiswa mendapat pengalaman di luar kampus; (3) dosen berkegiatan di luar kampus; (4) praktisi mengajar di kampus; (5) hasil kerja dosen digunakan masyarakat; (6) prodi berkerjasama dengan DUDI; (7) mata kuliah metode praktik; serta (8) prodi terakreditasi unggul.",
    ],
    analisis_paragrafs=[
        "Berdasarkan hasil AMI dan data PD Dikti, lulusan Universitas Tulungagung memiliki tingkat penyerapan di dunia kerja yang tinggi, dengan 86,50% lulusan terserap di DUDI dalam waktu kurang dari 6 bulan. Rata-rata IPK lulusan adalah 3,48, dengan rata-rata masa studi 3,8 tahun dan rata-rata masa tunggu kerja 3,2 bulan. Sebanyak 18,20% lulusan menjadi wirausahawan.",
        "Untuk publikasi ilmiah, berdasarkan AMI Penelitian, jumlah publikasi mencapai 95,60% dari target 10 dokumen per prodi per tahun. HKI (Paten, Hak Cipta, Desain Industri) mencapai 93,40% dengan tren meningkat setiap tahun. Sitasi karya ilmiah mencapai 90,80% dengan tren meningkat. Adopsi industri atau pemuatan hasil penelitian mencapai 88,50%.",
        "Untuk PkM, publikasi ilmiah dari hasil PkM mencapai 94,50%, HKI/TTG 93,80% (meningkat dari 20% di 2023), dan adopsi hasil PkM oleh masyarakat 91,20%. Persentase mahasiswa yang mengikuti MBKM mencapai 68,40% (target ≥60%), dan kontribusi mahasiswa dalam SDGs mencapai 76,20% (target ≥70%).",
    ],
    ami_data_table=[
        ["1", "Lulusan terserap DUDI <6 bulan", "86,50%", "≥ 70%", "Tercapai"],
        ["2", "Lulusan wirausahawan", "18,20%", "≥ 10%", "Tercapai"],
        ["3", "Rata-rata IPK lulusan", "3,48", "≥ 3,00", "Tercapai"],
        ["4", "Masa studi lulusan", "3,8 thn", "≤ 4,5 thn", "Tercapai"],
        ["5", "Masa tunggu kerja lulusan", "3,2 bln", "≤ 6 bln", "Tercapai"],
        ["6", "Publikasi ilmiah penelitian", "95,60%", "10 dok/prodi", "Tercapai"],
        ["7", "HKI penelitian", "93,40%", "Meningkat", "Tercapai"],
        ["8", "Sitasi karya ilmiah", "90,80%", "Meningkat", "Tercapai"],
        ["9", "Adopsi industri", "88,50%", "Ada/thn", "Tercapai"],
        ["10", "Publikasi PkM", "94,50%", "Meningkat", "Tercapai"],
        ["11", "HKI/TTG PkM", "93,80%", "Meningkat", "Tercapai"],
        ["12", "Adopsi hasil PkM masyarakat", "91,20%", "Ada/thn", "Tercapai"],
        ["13", "Mahasiswa MBKM", "68,40%", "≥ 60%", "Tercapai"],
        ["14", "Kontribusi mahasiswa SDGs", "76,20%", "≥ 70%", "Tercapai"],
        ["15", "Pengakuan kompetensi lulusan DUDI", "82,40%", "≥ 80%", "Tercapai"],
    ],
    kekuatan_list=[
        "86,50% lulusan terserap DUDI <6 bulan (target ≥70%)",
        "18,20% lulusan wirausahawan (target ≥10%)",
        "IPK rata-rata 3,48, masa studi 3,8 tahun, masa tunggu 3,2 bulan",
        "Publikasi penelitian 95,60% (10 dokumen/prodi/tahun)",
        "HKI/TTG PkM meningkat dari 20% (2023) menjadi 93,80% (2025)",
        "68,40% mahasiswa mengikuti MBKM (target ≥60%)",
        "8 IKU dengan baseline 2025 dan target 2030 telah ditetapkan",
    ],
    kelemahan_list=[
        "Adopsi industri masih 88,50%, perlu peningkatan",
        "Pengakuan kompetensi lulusan oleh DUDI 82,40%, perlu peningkatan",
        "Persentase prodi terakreditasi unggul masih 30% (target 70% pada 2030)",
        "Mata kuliah dengan metode praktik baru 45% (target 60% pada 2030)",
    ],
    rekomendasi_list=[
        "Meningkatkan kerja sama DUDI untuk penyerapan lulusan dan adopsi hasil penelitian",
        "Mendorong lebih banyak program studi untuk mengajukan akreditasi unggul",
        "Meningkatkan persentase mata kuliah dengan metode praktik dari 45% ke 60%",
        "Mengembangkan tracer study yang komprehensif untuk monitoring lulusan",
        "Meningkatkan pengakuan kompetensi lulusan melalui sertifikasi profesi",
        "Mengembangkan dashboard IKU yang terintegrasi dengan PD Dikti untuk monitoring real-time",
    ],
    ami_rata_skor="Berdasarkan AMI Pendidikan + Penelitian + PkM",
    ami_status="Semua indikator utama Tercapai",
    ami_dokumen="AMI Pendidikan + AMI Penelitian + AMI Pengabdian 2024-2025"
)

# ============================================================
# BAB XII - PENUTUP
# ============================================================

add_heading_custom(doc, "BAB XII\nPENUTUP", level=1)

add_heading_custom(doc, "12.1 Ringkasan Hasil Evaluasi Diri", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri yang mencakup 9 kriteria IAPT 4.1, Universitas Tulungagung menunjukkan kinerja yang baik dalam penyelenggaraan Tridharma Perguruan Tinggi. Dari total 212 indikator yang dievaluasi melalui 13 dokumen AMI 2024/2025, sebanyak 194 indikator (91,5%) berstatus Tercapai dan 18 indikator (8,5%) berstatus Belum Tercapai.")

add_body(doc, "Ringkasan hasil per kriteria adalah sebagai berikut:")
add_table_data(doc,
    ["Kriteria", "Dokumen AMI", "Indikator", "Tercapai", "Belum", "Persentase"],
    [
        ["1. VMTS", "AMI VMTS", "15", "10", "5", "66,7%"],
        ["2. Tata Pamong", "AMI TPTK + MBKM", "15", "12", "3", "80,0%"],
        ["3. Mahasiswa", "AMI Mahasiswa", "14", "14", "0", "100%"],
        ["4. SDM", "AMI SDM + Kesejahteraan", "18", "16", "2", "88,9%"],
        ["5. Keuangan/Sarpras", "AMI Keuangan + Sarpras", "14", "14", "0", "100%"],
        ["6. Pendidikan", "AMI Pendidikan + Proses", "19", "19", "0", "100%"],
        ["7. Penelitian", "AMI Penelitian", "20", "20", "0", "100%"],
        ["8. PkM", "AMI Pengabdian", "20", "20", "0", "100%"],
        ["9. Luaran Tridharma", "AMI Pendidikan+Pen+PkM", "15", "15", "0", "100%"],
        ["10. Pelayanan Mhs", "AMI Pelayanan", "20", "13", "7", "65,0%"],
        ["11. Proses Pembelajaran", "AMI Proses Pembelajaran", "20", "20", "0", "100%"],
        ["12. MBKM", "AMI MBKM", "4", "4", "0", "100%"],
        ["", "TOTAL", "212", "194", "18", "91,5%"],
    ],
    col_widths=[3, 3.5, 1.5, 1.5, 1.5, 2]
)

add_heading_custom(doc, "12.2 Rencana Tindak Lanjut", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri, Universitas Tulungagung menyusun Rencana Tindak Lanjut (RTL) dengan prioritas pada 18 indikator yang berstatus Belum Tercapai. RTL disusun dengan target waktu yang jelas dan penanggung jawab yang ditetapkan, sebagai berikut:")

add_table_data(doc,
    ["No", "Indikator Belum Tercapai", "RTL", "Target", "PJ"],
    [
        ["1", "Sosialisasi VMTS berkala (3,42)", "Diversifikasi media sosialisasi", "2026", "PPM"],
        ["2", "Metode sosialisasi VMTS (3,38)", "Handbook + media sosial", "2026", "PPM"],
        ["3", "Keterlibatan dosen evaluasi VMTS (3,36)", "Mekanisme feedback dosen", "2026", "Senat"],
        ["4", "Relevansi VMTS dunia kerja (3,48)", "Melibatkan DUDI dalam review", "2026", "Senat"],
        ["5", "Program mendukung VMTS (3,40)", "Penyelarasan program kerja", "2026", "Rektor"],
        ["6", "Dosen dalam evaluasi kebijakan (3,48)", "Forum dosen berkala", "2026", "Senat"],
        ["7", "Tata kelola dukung Tridharma (3,48)", "Penguatan koordinasi", "2026", "Rektor"],
        ["8", "Pimpinan terbuka masukan (3,49)", "Open door policy", "2026", "Rektor"],
        ["9", "Jabatan Lektor+ (52,40%)", "Program promosi jabatan", "2027", "BAKU"],
        ["10", "Sertifikat pendidik (94,60%)", "Pelatihan Serdos", "2027", "BAKU"],
        ["11-17", "Pelayanan mahasiswa (7 indikator)", "Perbaikan fasilitas & layanan", "2026", "BAA"],
        ["18", "Kerjasama implementasi", "AMI khusus Kerjasama", "2026", "BAU"],
    ],
    col_widths=[1, 4, 4, 1.5, 2]
)

add_heading_custom(doc, "12.3 Penutup", level=2)

add_body(doc, "Laporan Evaluasi Diri ini merupakan refleksi jujur dan komprehensif terhadap kinerja Universitas Tulungagung dalam penyelenggaraan Tridharma Perguruan Tinggi. Hasil evaluasi menunjukkan bahwa Universitas Tulungagung telah mencapai 91,5% dari total 212 indikator yang dievaluasi, dengan 8 kriteria mencapai 100% tercapai dan 4 kriteria masih memiliki indikator yang perlu diperbaiki.")

add_body(doc, "Universitas Tulungagung berkomitmen untuk terus meningkatkan mutu secara berkelanjutan melalui siklus PPEPP (Penetapan, Pelaksanaan, Evaluasi, Pengendalian, dan Peningkatan) sebagaimana diatur dalam Pasal 68 Permen 39 Tahun 2025. Rencana Tindak Lanjut yang telah disusun akan diimplementasikan dengan target waktu yang jelas dan penanggung jawab yang ditetapkan, dengan monitoring melalui AMI dan RTM tahunan.")

add_body(doc, "Dengan komitmen yang kuat dari seluruh sivitas akademika, sistem SPMI yang terstruktur dan terdokumentasi, serta dukungan data dari 13 dokumen AMI 2024/2025, Universitas Tulungagung siap untuk menghadapi proses akreditasi perguruan tinggi sesuai IAPT 4.1. Hasil evaluasi diri ini juga menjadi dasar bagi penyusunan strategi peningkatan mutu menuju tercapainya visi Universitas Tulungagung sebagai perguruan tinggi yang berkualitas dan mampu berkompetisi berskala nasional dan internasional.")

add_page_break(doc)

# ============================================================
# LAMPIRAN
# ============================================================

add_heading_custom(doc, "LAMPIRAN", level=1)

add_heading_custom(doc, "Lampiran 1: Daftar Dokumen AMI 2024/2025", level=2)

add_table_data(doc,
    ["No", "Dokumen AMI", "Kode Dokumen", "Tanggal", "Indikator"],
    [
        ["1", "AMI Pendidikan", "SPMI/PPM/MONEV/PEND/2025", "2024-2025", "19"],
        ["2", "AMI Penelitian", "SPMI/PPM/MONEV/PEN/2025", "2024-2025", "20"],
        ["3", "AMI Pengabdian Masyarakat", "SPMI/PPM/MONEV/PKM/2025", "2024-2025", "20"],
        ["4", "AMI VMTS", "SPMI/PPM/MONEV/VMTS/2025", "2024-2025", "15"],
        ["5", "AMI TPTK", "SPMI/PPM/MONEV/TPTK/2025", "2024-2025", "15"],
        ["6", "AMI Mahasiswa", "SPMI/PPM/MONEV/MHS/2025", "5 Maret 2025", "14"],
        ["7", "AMI SDM", "SPMI/PPM/MONEV/SDM/2025", "2024-2025", "18"],
        ["8", "AMI Kesejahteraan", "SPMI/PPM/MONEV/KSJH/2025", "24 Mei 2025", "19"],
        ["9", "AMI Keuangan", "SPMI/PPM/MONEV/KEU/2025", "15 Mei 2025", "14"],
        ["10", "AMI Sarana Prasarana", "SPMI/PPM/MONEV/SARPRAS/2025", "2024-2025", "14"],
        ["11", "AMI MBKM", "SPMI/PPM/MONEV/MBKM/2025", "5 Feb 2025", "4"],
        ["12", "AMI Pelayanan Mahasiswa", "SPMI/PPM/MONEV/PLYN/2025", "Ganjil & Genap", "20"],
        ["13", "AMI Proses Pembelajaran", "SPMI/PPM/MONEV/SPB/2025", "Ganjil & Genap", "20"],
    ],
    col_widths=[1, 3.5, 4, 2.5, 1.5]
)

add_heading_custom(doc, "Lampiran 2: Rekapitulasi 33 Standar SPMI", level=2)

add_body(doc, "Universitas Tulungagung memiliki 33 Standar SPMI yang terdiri dari 8 Standar Pendidikan, 8 Standar Penelitian, 8 Standar Pengabdian kepada Masyarakat, dan 9 Standar Tambahan sebagai pelampauan SN Dikti. Setiap standar memiliki kode dokumen unik dan dikelola melalui siklus PPEPP.")

add_table_data(doc,
    ["No", "Nama Standar", "Kode", "Kategori", "Skor AMI 2025"],
    [
        ["1", "Standar Kompetensi Lulusan", "SKL", "Pendidikan", "87"],
        ["2", "Standar Proses Pembelajaran", "SPB", "Pendidikan", "89"],
        ["3", "Standar Penilaian", "SPN", "Pendidikan", "85"],
        ["4", "Standar Pengelolaan", "SPP", "Pendidikan", "84"],
        ["5", "Standar Isi", "SIS", "Pendidikan", "83"],
        ["6", "Standar Dosen & Tendik", "SDT", "Pendidikan", "89"],
        ["7", "Standar Sarana & Prasarana", "SSP", "Pendidikan", "95"],
        ["8", "Standar Pembiayaan", "SPM", "Pendidikan", "95"],
        ["9", "Standar Luaran Penelitian", "PEN-LRN", "Penelitian", "94"],
        ["10", "Standar Proses Penelitian", "PEN-PRS", "Penelitian", "95"],
        ["11", "Standar Masukan Penelitian", "PEN-MSK", "Penelitian", "91"],
        ["12", "Standar Hasil Penelitian", "PEN-HSL", "Penelitian", "92"],
        ["13", "Standar Publikasi Penelitian", "PEN-PUB", "Penelitian", "94"],
        ["14", "Standar Paten & HKI", "PEN-PTN", "Penelitian", "93"],
        ["15", "Standar Diseminasi Penelitian", "PEN-DSM", "Penelitian", "90"],
        ["16", "Standar Kolaborasi Penelitian", "PEN-KLB", "Penelitian", "89"],
        ["17-24", "Standar PkM (8 standar)", "PKM-*", "PkM", "89-96"],
        ["25", "Standar VMTS", "VMTS", "Tambahan", "89"],
        ["26", "Standar TPTK", "TPTK", "Tambahan", "90"],
        ["27", "Standar Mahasiswa", "MHS", "Tambahan", "93"],
        ["28", "Standar SDM", "SDM", "Tambahan", "89"],
        ["29", "Standar Keuangan/Sarpras", "KSP", "Tambahan", "95"],
        ["30", "Standar Kerjasama", "KJS", "Tambahan", "80"],
        ["31", "Standar MBKM", "MBKM", "Tambahan", "96"],
        ["32", "Standar Luaran Tridharma", "LCT", "Tambahan", "88"],
        ["33", "Standar Pelayanan Mhs", "PKM-MHS", "Tambahan", "88"],
    ],
    col_widths=[1.5, 4, 2, 2.5, 2]
)

add_heading_custom(doc, "Lampiran 3: 8 IKU dan Target Capaian", level=2)

add_table_data(doc,
    ["No", "IKU", "Baseline 2025", "Target 2030", "Tren"],
    [
        ["1", "Lulusan terserap DUDI <6 bulan", "87%", "92%", "↑"],
        ["2", "Mahasiswa pengalaman luar kampus", "68%", "40%", "↑"],
        ["3", "Dosen berkegiatan luar kampus", "15%", "35%", "↑"],
        ["4", "Praktisi mengajar di kampus", "2/prodi", "6/prodi", "↑"],
        ["5", "Hasil kerja dosen digunakan masyarakat", "18", "35", "↑"],
        ["6", "Prodi berkerjasama DUDI", "3/prodi", "7/prodi", "↑"],
        ["7", "Mata kuliah metode praktik", "45%", "60%", "↑"],
        ["8", "Prodi terakreditasi unggul", "30%", "70%", "↑"],
    ],
    col_widths=[1, 5, 2.5, 2.5, 1.5]
)

add_spacer(doc, 3)

# Final note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- END OF REPORT ---")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nLaporan Evaluasi Diri Universitas Tulungagung\nAkreditasi Perguruan Tinggi (APT) 4.1\nTahun Akademik 2024/2025\n\nPusat Penjaminan Mutu (PPM)\nUniversitas Tulungagung\nSK Rektor Nomor: A/002.I/KEP/UNITA/I/2025")
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

# ============================================================
# SAVE DOCUMENT
# ============================================================

output_path = "/home/z/my-project/download/LAPORAN_EVALUASI_DIRI_APT_4_1_UNITA_2025.docx"

# ============================================================
# TAMBAHAN: SWOT ANALYSIS, BENCHMARKING, DETAIL EXPANSION
# ============================================================

add_page_break(doc)

# BAB XIII - ANALISIS SWOT MENYELURUH
add_heading_custom(doc, "BAB XIII\nANALISIS SWOT UNIVERSITAS TULUNGAGUNG", level=1)

add_heading_custom(doc, "13.1 Strengths (Kekuatan)", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri yang komprehensif melalui 13 dokumen AMI 2024/2025, Universitas Tulungagung memiliki sejumlah kekuatan yang menjadi modal strategis untuk pengembangan institusi ke depan. Kekuatan-kekuatan ini teridentifikasi dari capaian indikator yang berstatus Tercapai dengan skor tinggi, serta praktik baik yang telah diimplementasikan secara konsisten.")

add_body(doc, "Pertama, sistem penjaminan mutu internal (SPMI) yang terstruktur dan terdokumentasi dengan empat dokumen mutu lengkap (Kebijakan Mutu 65 halaman, Manual Mutu 55 halaman, Standar Mutu 312 halaman, Formulir Mutu 132 halaman) yang seluruhnya telah diselaraskan dengan Permen 39 Tahun 2025. Sistem SPMI ini didukung oleh siklus PPEPP yang konsisten dan struktur organisasi PPM-LPM-GKM yang jelas di setiap aras.")

add_body(doc, "Kedua, capaian lulusan yang sangat baik dengan tingkat penyerapan di dunia kerja mencapai 86,50% (di atas target 70%), IPK rata-rata 3,48, masa studi 3,8 tahun (di bawah target 4,5 tahun), dan masa tunggu kerja 3,2 bulan (jauh di bawah target 6 bulan). Sebanyak 18,20% lulusan menjadi wirausahawan, menunjukkan bahwa UNITA tidak hanya menghasilkan lulusan yang siap kerja tetapi juga mampu menciptakan lapangan kerja.")

add_body(doc, "Ketiga, pengelolaan keuangan yang sehat dengan realisasi pendapatan 95,20% dari target RENOP, rasio pendapatan terhadap belanja 97,30%, dan seluruh 14 indikator AMI Keuangan berstatus Tercapai. Total pendapatan target mencapai Rp28.400 juta dengan diversifikasi sumber pendapatan dari pemerintah, mahasiswa, kegiatan profesional, dan sumber lainnya.")

add_body(doc, "Keempat, sarana dan prasarana yang meningkat signifikan, dengan sarana baik mencapai 95,20% (dari 83% di 2023) dan prasarana baik 95,80% (dari 84%). Seluruh 14 indikator AMI Sarana Prasarana berstatus Tercapai, mencakup ruang kelas, laboratorium, perpustakaan, fasilitas ICT, dan fasilitas kemahasiswaan.")

add_body(doc, "Kelima, kinerja penelitian dan PkM yang sangat baik dengan seluruh 20 indikator AMI Penelitian dan 20 indikator AMI PkM berstatus Tercapai. Publikasi ilmiah mencapai 95,60%, HKI 93,40%, dan TTG PkM meningkat dari 20% (2023) menjadi 93,80% (2025). Adopsi hasil PkM oleh masyarakat mencapai 91,20%.")

add_body(doc, "Keenam, kesejahteraan SDM yang sangat baik dengan rata-rata 96,32% dari 25 indikator, seluruhnya Tercapai. Ini mencakup buku pedoman, SOP, hak karyawan, imbal jasa, asuransi, cuti, dana pensiun, program rekreasi, dan survei kepuasan. Tingkat kepuasan karyawan terhadap kesejahteraan mencapai 94,80%.")

add_body(doc, "Ketujuh, pelaksanaan MBKM yang sangat baik dengan rata-rata 95,50 (Sangat Baik) dari 4 kelompok indikator. Dokumen Standar MBKM dan Buku Panduan mencapai skor 100, pemetaan mitra 94,20, kerjasama mitra 92,40, dan rekognisi/konversi nilai 96,90. Partisipasi mahasiswa MBKM mencapai 68,40% (target 60%).")

add_heading_custom(doc, "13.2 Weaknesses (Kelemahan)", level=2)

add_body(doc, "Meskipun memiliki banyak kekuatan, Universitas Tulungagung juga mengidentifikasi sejumlah kelemahan yang perlu diatasi untuk peningkatan mutu berkelanjutan. Kelemahan-kelemahan ini teridentifikasi dari 18 indikator yang berstatus Belum Tercapai dari total 212 indikator AMI.")

add_body(doc, "Pertama, pada Kriteria 1 (VMTS), terdapat 5 indikator yang belum tercapai, terkait dengan sosialisasi VMTS yang belum optimal (skor 3,38-3,42), keterlibatan dosen dalam evaluasi VMTS yang rendah (skor 3,36), dan relevansi VMTS dengan kebutuhan dunia kerja (skor 3,48). Hal ini menunjukkan bahwa VMTS belum sepenuhnya dipahami dan diinternalisasi oleh seluruh sivitas akademika.")

add_body(doc, "Kedua, pada Kriteria 2 (Tata Pamong), terdapat 3 indikator yang belum tercapai, terkait dengan keterlibatan dosen dalam penyusunan kebijakan (skor 3,48), tata kelola yang mendukung Tridharma (skor 3,48), dan keterbukaan pimpinan terhadap masukan (skor 3,49). Ketiga indikator ini berada sedikit di bawah target 3,50 namun menunjukkan area yang perlu perbaikan.")

add_body(doc, "Ketiga, pada Kriteria 4 (SDM), terdapat 2 indikator yang belum tercapai: persentase dosen dengan jabatan fungsional minimal Lektor (52,40% dari target 75%) dan kepemilikan sertifikat pendidik (94,60% dari target 100%). Kedua indikator ini memerlukan waktu dan investasi yang signifikan untuk diperbaiki.")

add_body(doc, "Keempat, pada Kriteria 10 (Pelayanan Mahasiswa), terdapat 7 indikator yang belum tercapai, mencakup kecepatan respons terhadap keluhan (3,47), transparansi penggunaan dana (3,46), kecepatan verifikasi dana (3,48), akses informasi organisasi mahasiswa (3,49), jadwal pembimbingan (3,48), pendampingan PKM (3,49), dan fasilitas kebersihan (3,44). Ini adalah area dengan jumlah kelemahan terbanyak.")

add_body(doc, "Kelima, pada Kriteria 2 (Kerjasama), tidak ada AMI khusus untuk kerja sama, sehingga evaluasi masih bersifat estimasi berdasarkan AMI MBKM. Implementasi kerja sama hanya mencapai 60% dari total MoU/MoA yang aktif, menunjukkan adanya gap antara penandatanganan kerja sama dan implementasi nyata.")

add_heading_custom(doc, "13.3 Opportunities (Peluang)", level=2)

add_body(doc, "Universitas Tulungagung memiliki sejumlah peluang strategis yang dapat dimanfaatkan untuk pengembangan institusi. Peluang-peluang ini berasal dari faktor eksternal yang dapat mendukung pencapaian visi dan misi universitas.")

add_body(doc, "Pertama, berlakunya Permen 39 Tahun 2025 yang membuka peluang penyederhanaan mekanisme akreditasi berbasis PD Dikti. Pasal 78 ayat (2) menyebutkan bahwa mekanisme perpanjangan status terakreditasi disusun dengan memperhatikan efektivitas, efisiensi, dan meminimumkan beban administratif perguruan tinggi dengan memanfaatkan data dari PD Dikti. Hal ini dapat mengurangi beban administratif dalam proses akreditasi.")

add_body(doc, "Kedua, tersedianya berbagai Lembaga Akreditasi Mandiri (LAM) yang memungkinkan akreditasi spesifik per rumpun ilmu. Universitas Tulungagung dapat memilih LAM yang paling sesuai dengan karakteristik setiap program studi, sehingga evaluasi akreditasi menjadi lebih relevan dan kontekstual.")

add_body(doc, "Ketiga, peluang memperoleh status terakreditasi unggul dan terakreditasi secara internasional sebagaimana diatur dalam Pasal 73-74 dan Pasal 107 Permen 39 Tahun 2025. Saat ini, 30% program studi UNITA telah berstatus terakreditasi unggul, dengan target 70% pada 2030.")

add_body(doc, "Keempat, kemajuan teknologi informasi yang memungkinkan otomasi proses SPMI. Universitas Tulungagung dapat mengembangkan dashboard SPMI berbasis IT yang terintegrasi dengan PD Dikti, sehingga monitoring dan evaluasi dapat dilakukan secara real-time.")

add_body(doc, "Kelima, kebijakan Merdeka Belajar Kampus Merdeka (MBKM) yang membuka peluang pengayaan pengalaman mahasiswa. Universitas Tulungagung telah mengembangkan 6 skema MBKM dengan 30 mitra DUDI, dan masih ada ruang untuk ekspansi lebih lanjut.")

add_body(doc, "Keenam, meningkatnya kesadaran masyarakat akan mutu pendidikan tinggi yang membuka peluang kerja sama strategis dengan DUDI. Universitas Tulungagung dapat memanfaatkan ini untuk meningkatkan kerja sama dalam bidang penelitian terapan, magang, dan penyerapan lulusan.")

add_heading_custom(doc, "13.4 Threats (Tantangan)", level=2)

add_body(doc, "Selain peluang, Universitas Tulungagung juga menghadapi sejumlah tantangan yang perlu diantisipasi dan dimitigasi. Tantangan-tantangan ini berasal dari faktor eksternal yang dapat menghambat pencapaian tujuan institusi.")

add_body(doc, "Pertama, persaingan antar perguruan tinggi yang semakin ketat, baik dari sisi kualitas maupun kuantitas. Universitas Tulungagung perlu terus berinovasi untuk mempertahankan dan meningkatkan daya saing, terutama dalam menarik calon mahasiswa baru dan mempertahankan retensi mahasiswa.")

add_body(doc, "Kedua, perubahan regulasi yang dinamis. Walaupun Permen 39 Tahun 2025 telah berlaku, masih ada kemungkinan perubahan atau penyesuaian regulasi di masa mendatang. Universitas Tulungagung perlu memastikan bahwa sistem SPMI yang adaptif dan responsif terhadap perubahan regulasi.")

add_body(doc, "Ketiga, keterbatasan sumber daya finansial sebagai perguruan tinggi swasta. Meskipun pengelolaan keuangan telah berjalan baik, Universitas Tulungagung masih perlu mendiversifikasi sumber pendapatan untuk mendukung investasi jangka panjang, terutama dalam pengembangan SDM dan sarana-prasarana.")

add_body(doc, "Keempat, persaingan dalam menarik dan mempertahankan dosen berkualitas. Universitas Tulungagung perlu menyediakan paket kompensasi dan benefit yang kompetitif, serta program pengembangan karir yang menarik, untuk dapat bersaing dengan perguruan tinggi negeri dan swasta lainnya.")

add_body(doc, "Kelima, ekspektasi pemangku kepentingan yang terus meningkat. Masyarakat, pemerintah, dan DUDI menuntut lulusan yang tidak hanya kompeten secara akademik, tetapi juga memiliki soft skills, digital literacy, dan kesiapan kerja. Universitas Tulungagung perlu terus menyesuaikan kurikulum dan metode pembelajaran untuk memenuhi ekspektasi ini.")

# SWOT Matrix Table
add_spacer(doc, 2)
p = doc.add_paragraph()
run = p.add_run("Tabel 13.1 Matriks SWOT Universitas Tulungagung")
run.font.bold = True
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
headers = ["", "Peluang (O)", "Tantangan (T)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, "1A2B5C")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)

# Row 1: SO Strategy
cell = table.rows[1].cells[0]
cell.text = "Kekuatan (S)"
set_cell_shading(cell, "10B981")
for p in cell.paragraphs:
    for run in p.runs:
        run.font.color.rgb = WHITE
        run.font.bold = True

cell = table.rows[1].cells[1]
cell.text = ("STRATEGI SO:\n"
    "1. Manfaatkan SPMI kuat untuk akreditasi unggul\n"
    "2. Gunakan data AMI untuk efisiensi akreditasi\n"
    "3. Kembangkan MBKM dengan DUDI\n"
    "4. Ekspansi LAM untuk prodi spesifik\n"
    "5. Manfaatkan teknologi untuk dashboard SPMI")
set_cell_shading(cell, "F0FDF4")

cell = table.rows[1].cells[2]
cell.text = ("STRATEGI ST:\n"
    "1. Diferensiasi melalui kualitas lulusan\n"
    "2. Adaptasi cepat terhadap regulasi\n"
    "3. Diversifikasi pendapatan via RGU\n"
    "4. Retensi dosen via kesejahteraan 96%\n"
    "5. Update kurikulum berbasis DUDI")
set_cell_shading(cell, "FEF3C7")

# Row 2: WO Strategy
cell = table.rows[2].cells[0]
cell.text = "Kelemahan (W)"
set_cell_shading(cell, "F59E0B")
for p in cell.paragraphs:
    for run in p.runs:
        run.font.color.rgb = WHITE
        run.font.bold = True

cell = table.rows[2].cells[1]
cell.text = ("STRATEGI WO:\n"
    "1. Perbaiki sosialisasi VMTS via teknologi\n"
    "2. Tingkatkan keterlibatan dosen via forum\n"
    "3. Promosi jabatan akademik dosen\n"
    "4. Sertifikasi pendidik dosen\n"
    "5. Perbaiki pelayanan via One-Stop Service")
set_cell_shading(cell, "FFFBEB")

cell = table.rows[2].cells[2]
cell.text = ("STRATEGI WT:\n"
    "1. Fokus perbaikan 18 indikator belum tercapai\n"
    "2. AMI khusus Kerjasama\n"
    "3. Investasi smart campus\n"
    "4. Program promosi Lektor massal\n"
    "5. Perbaikan fasilitas kebersihan")
set_cell_shading(cell, "FEE2E2")

add_page_break(doc)

# ============================================================
# BAB XIV - BENCHMARKING & ANALISIS PERBANDINGAN
# ============================================================

add_heading_custom(doc, "BAB XIV\nBENCHMARKING DAN ANALISIS PERBANDINGAN", level=1)

add_heading_custom(doc, "14.1 Perbandingan Capaian AMI Antar Kriteria", level=2)

add_body(doc, "Analisis perbandingan capaian AMI antar kriteria memberikan gambaran tentang area-area yang telah berjalan dengan baik dan area yang memerlukan perhatian khusus. Berikut adalah perbandingan persentase ketercapaian indikator AMI per kriteria:")

add_table_data(doc,
    ["Kriteria", "Total Indikator", "Tercapai", "Belum", "Persentase", "Kategori"],
    [
        ["1. VMTS", "15", "10", "5", "66,7%", "Baik"],
        ["2. Tata Pamong", "15", "12", "3", "80,0%", "Baik"],
        ["3. Mahasiswa", "14", "14", "0", "100%", "Sangat Baik"],
        ["4. SDM", "18", "16", "2", "88,9%", "Sangat Baik"],
        ["5. Keuangan/Sarpras", "14", "14", "0", "100%", "Sangat Baik"],
        ["6. Pendidikan", "19", "19", "0", "100%", "Sangat Baik"],
        ["7. Penelitian", "20", "20", "0", "100%", "Sangat Baik"],
        ["8. PkM", "20", "20", "0", "100%", "Sangat Baik"],
        ["9. Luaran Tridharma", "15", "15", "0", "100%", "Sangat Baik"],
        ["10. Pelayanan Mhs", "20", "13", "7", "65,0%", "Baik"],
        ["11. Proses Pembelajaran", "20", "20", "0", "100%", "Sangat Baik"],
        ["12. MBKM", "4", "4", "0", "100%", "Sangat Baik"],
        ["13. Kesejahteraan", "18", "18", "0", "100%", "Sangat Baik"],
        ["TOTAL", "212", "194", "18", "91,5%", "Sangat Baik"],
    ],
    col_widths=[3, 2, 1.5, 1.5, 1.5, 2.5]
)

add_body(doc, "Dari tabel di atas, terlihat bahwa 8 kriteria mencapai 100% ketercapaian indikator (Mahasiswa, Keuangan/Sarpras, Pendidikan, Penelitian, PkM, Luaran Tridharma, Proses Pembelajaran, MBKM, dan Kesejahteraan). Kriteria dengan persentase terendah adalah Pelayanan kepada Mahasiswa (65,0%) dan VMTS (66,7%), yang menjadi prioritas perbaikan.")

add_heading_custom(doc, "14.2 Tren Capaian Proses Pembelajaran", level=2)

add_body(doc, "AMI Proses Pembelajaran dilaksanakan dua kali per tahun (Ganjil dan Genap), memungkinkan analisis tren yang lebih detail. Berikut adalah perbandingan skor rata-rata antar periode:")

add_table_data(doc,
    ["Periode", "Skor Rata-rata", "Kategori", "Tren"],
    [
        ["Ganjil 2023-2", "3,43", "Baik", "Baseline"],
        ["Ganjil 2024-1", "3,49", "Baik", "↑ +0,06"],
        ["Genap 2024-2", "3,59", "Sangat Baik", "↑ +0,10"],
        ["Rata-rata 2024/2025", "3,54", "Sangat Baik", "↑ +0,11"],
    ],
    col_widths=[4, 3, 3, 3]
)

add_body(doc, "Tren peningkatan skor dari 3,43 menjadi 3,59 menunjukkan efektivitas implementasi Rencana Tindak Lanjut (RTL) dari periode sebelumnya. Peningkatan ini terjadi di seluruh fakultas (FE, FH, FISIP, FP, FT, F.KES), dengan praktik baik dari Prodi Akuntansi yang dapat direplikasi ke prodi lain.")

add_heading_custom(doc, "14.3 Tren Capaian Pelayanan Mahasiswa", level=2)

add_body(doc, "AMI Pelayanan kepada Mahasiswa juga dilaksanakan dua kali per tahun. Berikut adalah perbandingan skor:")

add_table_data(doc,
    ["Periode", "Skor Rata-rata", "Kategori", "Indikator ≥3,50"],
    [
        ["Ganjil 2023-1", "3,42", "Baik", "8 dari 20"],
        ["Ganjil 2024-1", "3,48", "Baik", "11 dari 20"],
        ["Genap 2024-2", "3,58", "Sangat Baik", "17 dari 20"],
    ],
    col_widths=[4, 3, 3, 3]
)

add_body(doc, "Peningkatan dari 3,42 menjadi 3,58 dan dari 8 menjadi 17 indikator yang mencapai target menunjukkan kemajuan yang signifikan. Namun, 3 indikator masih di bawah target 3,50 dan menjadi prioritas perbaikan untuk periode berikutnya.")

add_heading_custom(doc, "14.4 Analisis Capaian IKU", level=2)

add_body(doc, "Delapan Indikator Kinerja Utama (IKU) Universitas Tulungagung dengan baseline 2025 dan target 2030:")

add_table_data(doc,
    ["IKU", "Baseline 2025", "Capaian 2025", "Target 2030", "Progress"],
    [
        ["1. Lulusan terserap DUDI <6 bln", "60%", "87%", "92%", "95%"],
        ["2. Mahasiswa pengalaman luar", "20%", "68%", "40%", "170%"],
        ["3. Dosen di luar kampus", "15%", "15%", "35%", "43%"],
        ["4. Praktisi mengajar", "2/prodi", "2/prodi", "6/prodi", "33%"],
        ["5. Hasil dosen digunakan", "15", "18", "35", "51%"],
        ["6. Prodi kerjasama DUDI", "3/prodi", "3/prodi", "7/prodi", "43%"],
        ["7. Mata kuliah metode praktik", "40%", "45%", "60%", "75%"],
        ["8. Prodi akreditasi unggul", "30%", "30%", "70%", "43%"],
    ],
    col_widths=[4, 2, 2, 2, 2]
)

add_body(doc, "Dari tabel di atas, IKU 1 (lulusan terserap) dan IKU 2 (mahasiswa pengalaman luar) telah melampaui target 2030. IKU 7 (mata kuliah praktik) berada pada 75% progress. IKU 3, 4, 5, 6, dan 8 masih berada di bawah 50% progress dan memerlukan akselerasi.")

add_page_break(doc)

# ============================================================
# BAB XV - ANALISIS RISIKO DAN MITIGASI
# ============================================================

add_heading_custom(doc, "BAB XV\nANALISIS RISIKO DAN MITIGASI", level=1)

add_heading_custom(doc, "15.1 Identifikasi Risiko", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri, Universitas Tulungagung mengidentifikasi empat jenis risiko utama dalam pelaksanaan SPMI, sebagaimana diatur dalam Manual Mutu SPMI UNITA 2025. Keempat jenis risiko tersebut adalah risiko kebijakan, risiko kepatuhan, risiko operasional, dan risiko reputasi.")

add_table_data(doc,
    ["Jenis Risiko", "Dampak Potensial", "Tingkat", "Mitigasi"],
    [
        ["Risiko Kebijakan", "Kebijakan SPMI tidak selaras dengan regulasi", "Sedang", "Kaji berkala regulasi; konsultasi Senat"],
        ["Risiko Kepatuhan", "Data PD Dikti tidak tepat waktu; temuan AMI tidak ditindaklanjuti", "Tinggi", "Monitoring berkala; sanksi; eskalasi"],
        ["Risiko Operasional", "Sistem down; auditor tidak tersedia; dokumen hilang", "Sedang", "SOP; pelatihan; sistem back-up"],
        ["Risiko Reputasi", "Penurunan akreditasi; publikasi negatif", "Rendah", "Tindak lanjut cepat; transparansi; PR aktif"],
    ],
    col_widths=[3, 4, 1.5, 4]
)

add_heading_custom(doc, "15.2 Risiko Spesifik per Kriteria", level=2)

add_body(doc, "Setiap kriteria memiliki risiko spesifik yang perlu dimitigasi. Berikut adalah analisis risiko per kriteria:")

add_table_data(doc,
    ["Kriteria", "Risiko Utama", "Dampak", "Mitigasi"],
    [
        ["1. VMTS", "VMTS tidak dipahami sivitas", "Penurunan kinerja", "Diversifikasi sosialisasi"],
        ["2. Tata Pamong", "Kebijakan tidak diimplementasikan", "Inefisiensi", "Monitoring implementasi"],
        ["3. Mahasiswa", "Penurunan jumlah mahasiswa", "Penurunan pendapatan", "Strategi pemasaran; beasiswa"],
        ["4. SDM", "Dosen resign; sertifikat tidak tercapai", "Penurunan kualitas", "Retensi; pelatihan Serdos"],
        ["5. Keuangan", "Opini audit tidak WTP", "Penurunan kepercayaan", "Tata kelola keuangan"],
        ["6. Pendidikan", "Kurikulum tidak relevan", "Lulusan tidak terserap", "Review kurikulum; tracer study"],
        ["7. Penelitian", "Publikasi menurun", "Penurunan reputasi", "Insentif; pendanaan"],
        ["8. PkM", "Adopsi hasil rendah", "Dampak masyarakat rendah", "Diseminasi; kerjasama DUDI"],
        ["9. Luaran", "IKU tidak tercapai", "Akreditasi turun", "Dashboard monitoring; RTL"],
    ],
    col_widths=[2.5, 3.5, 2.5, 4]
)

add_heading_custom(doc, "15.3 Strategi Mitigasi", level=2)

add_body(doc, "Strategi mitigasi risiko di Universitas Tulungagung dilaksanakan melalui beberapa mekanisme: (1) monitoring berkala oleh PPM, LPM, dan GKM terhadap pelaksanaan standar; (2) Audit Mutu Internal (AMI) tahunan untuk evaluasi sistemik; (3) Rapat Tinjauan Manajemen (RTM) tahunan untuk evaluasi formal manajemen puncak; (4) sistem informasi SPMI berbasis IT untuk monitoring real-time; serta (5) benchmarking dengan perguruan tinggi lain untuk pembelajaran dan perbaikan.")

add_body(doc, "Selain itu, Universitas Tulungagung juga menerapkan prinsip triangulasi data sebagaimana diatur dalam Pasal 66 Permen 39 Tahun 2025, yang memastikan bahwa setiap keputusan strategis berbasis pada data yang valid dari minimal tiga sumber. Hal ini mengurangi risiko pengambilan keputusan yang berbasis pada data yang tidak akurat atau tidak lengkap.")

add_page_break(doc)

# ============================================================
# BAB XVI - RENCANA STRATEGIS PENINGKATAN MUTU
# ============================================================

add_heading_custom(doc, "BAB XVI\nRENCANA STRATEGIS PENINGKATAN MUTU 2025-2030", level=1)

add_heading_custom(doc, "16.1 Visi dan Misi Mutu", level=2)

add_body(doc, "Visi mutu Universitas Tulungagung adalah menjadi perguruan tinggi yang berkualitas dan mampu berkompetisi berskala nasional dan internasional, dengan sistem penjaminan mutu internal yang kuat, transparan, dan akuntabel. Misi mutu adalah memastikan pemenuhan dan pelampauan SN Dikti secara sistematis dan berkelanjutan melalui siklus PPEPP.")

add_heading_custom(doc, "16.2 Target Strategis 2025-2030", level=2)

add_table_data(doc,
    ["Aspek", "Baseline 2025", "Target 2026", "Target 2028", "Target 2030"],
    [
        ["Rata-rata Skor AMI", "91/100", "93/100", "96/100", "98/100"],
        ["Indikator Tercapai", "91,5%", "95%", "98%", "100%"],
        ["Lulusan terserap DUDI", "87%", "88%", "90%", "92%"],
        ["Mahasiswa MBKM", "68%", "75%", "82%", "90%"],
        ["Prodi akreditasi unggul", "30%", "40%", "55%", "70%"],
        ["Publikasi Scopus", "20", "35", "68", "110"],
        ["Paten granted", "1", "2", "7", "15"],
        ["Dosen Serdos", "94,6%", "96%", "98%", "100%"],
        ["Dosen Lektor+", "52,4%", "65%", "80%", "90%"],
        ["Opini audit", "WDP", "WTP", "WTP", "WTP"],
        ["MoU aktif", "40", "55", "100", "165"],
        ["Kerjasama internasional", "10", "18", "35", "62"],
    ],
    col_widths=[4, 2.5, 2.5, 2.5, 2.5]
)

add_heading_custom(doc, "16.3 Program Prioritas", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri, Universitas Tulungagung menetapkan program prioritas sebagai berikut:")

programs = [
    ("Program 1: Penguatan VMTS", "Diversifikasi media sosialisasi VMTS, melibatkan DUDI dalam review VMTS, mengembangkan dashboard monitoring VMTS real-time. Target: pemahaman sivitas ≥90% pada 2026."),
    ("Program 2: Penguatan Tata Kelola", "Meningkatkan keterlibatan dosen dalam penyusunan kebijakan, penguatan koordinasi antar-unit, open door policy pimpinan. Target: semua indikator TPTS ≥3,50 pada 2026."),
    ("Program 3: Pengembangan SDM", "Program promosi jabatan akademik massal, pelatihan Serdos bagi dosen yang belum bersertifikat, studi lanjut S3 bagi dosen. Target: Lektor+ ≥65% dan Serdos ≥96% pada 2026."),
    ("Program 4: Perbaikan Pelayanan Mahasiswa", "Implementasi One-Stop Service, perbaikan fasilitas kebersihan, peningkatan kecepatan respons, transparansi penggunaan dana. Target: 20/20 indikator ≥3,50 pada 2026."),
    ("Program 5: Pengembangan Akreditasi", "Mendorong program studi untuk mengajukan akreditasi unggul, benchmarking dengan PT unggulan, persiapan instrumen akreditasi terbaru. Target: 40% prodi unggul pada 2026."),
    ("Program 6: Penguatan Penelitian & PkM", "Meningkatkan publikasi Scopus, pengajuan paten, kerja sama DUDI untuk penelitian terapan, komersialisasi hasil PkM. Target: publikasi Scopus 35 dan paten 2 pada 2026."),
    ("Program 7: Pengembangan MBKM", "Ekspansi skema MBKM, peningkatan mitra DUDI, pengembangan joint program dengan PT mitra. Target: 75% mahasiswa mengikuti MBKM pada 2026."),
    ("Program 8: Smart Campus", "Pengembangan WiFi campus-wide, smart classroom, IoT, sistem informasi terintegrasi. Target: 10 smart campus components pada 2026."),
    ("Program 9: Diversifikasi Pendapatan", "Pengembangan RGU, kerja sama DUDI untuk pendanaan, hibah eksternal. Target: pendapatan RGU ≥9% pada 2026."),
    ("Program 10: AMI Khusus Kerjasama", "Menyelenggarakan AMI khusus untuk kerja sama, evaluasi implementasi MoU/MoA. Target: AMI Kerjasama dilaksanakan pada 2026."),
]

for title, desc in programs:
    add_heading_custom(doc, title, level=3)
    add_body(doc, desc)

add_heading_custom(doc, "16.4 Timeline Implementasi", level=2)

add_table_data(doc,
    ["Tahun", "Program Prioritas", "Target Capaian"],
    [
        ["2025", "Baseline data AMI, penyusunan RTL", "91,5% indikator tercapai"],
        ["2026", "Program 1-4 (VMTS, Tata Kelola, SDM, Pelayanan)", "95% indikator tercapai"],
        ["2027", "Program 5-7 (Akreditasi, Penelitian, MBKM)", "97% indikator tercapai"],
        ["2028", "Program 8-9 (Smart Campus, Diversifikasi)", "98% indikator tercapai"],
        ["2029", "Konsolidasi dan evaluasi menyeluruh", "99% indikator tercapai"],
        ["2030", "Pencapaian target strategis", "100% indikator tercapai"],
    ],
    col_widths=[2, 5, 4]
)

add_page_break(doc)

# ============================================================
# BAB XVII - KESIMPULAN DAN REKOMENDASI AKHIR
# ============================================================

add_heading_custom(doc, "BAB XVII\nKESIMPULAN DAN REKOMENDASI AKHIR", level=1)

add_heading_custom(doc, "17.1 Kesimpulan", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri yang komprehensif terhadap 9 kriteria IAPT 4.1, mencakup 212 indikator dari 13 dokumen AMI 2024/2025, dapat disimpulkan bahwa Universitas Tulungagung telah menunjukkan kinerja yang sangat baik dalam penyelenggaraan Tridharma Perguruan Tinggi. Dengan 194 indikator (91,5%) berstatus Tercapai, Universitas Tulungagung telah memenuhi standar mutu yang ditetapkan sesuai dengan Permen 39 Tahun 2025.")

add_body(doc, "Sembilan kriteria menunjukkan capaian yang bervariasi: 8 kriteria mencapai 100% ketercapaian (Mahasiswa, Keuangan/Sarpras, Pendidikan, Penelitian, PkM, Luaran Tridharma, Proses Pembelajaran, MBKM, Kesejahteraan), sementara 4 kriteria masih memiliki indikator yang perlu diperbaikan (VMTS 66,7%, Tata Pamong 80%, SDM 88,9%, Pelayanan Mahasiswa 65%). Area-area yang memerlukan perbaikan telah diidentifikasi dan rencana tindak lanjut telah disusun dengan target waktu dan penanggung jawab yang jelas.")

add_body(doc, "Sistem Penjaminan Mutu Internal (SPMI) Universitas Tulungagung telah berjalan dengan baik, didukung oleh empat dokumen mutu yang lengkap dan terdokumentasi (Kebijakan Mutu, Manual Mutu, Standar Mutu, Formulir Mutu), struktur organisasi PPM-LPM-GKM yang jelas, serta siklus PPEPP yang konsisten. Prinsip triangulasi data berbasis PD Dikti telah diimplementasikan dalam pengumpulan dan analisis data AMI.")

add_body(doc, "Capaian lulusan Universitas Tulungagung sangat membanggakan, dengan 86,50% lulusan terserap di DUDI dalam waktu kurang dari 6 bulan, IPK rata-rata 3,48, masa studi 3,8 tahun, dan 18,20% lulusan menjadi wirausahawan. Capaian ini melampaui target yang ditetapkan dan menunjukkan bahwa Universitas Tulungagung berhasil menghasilkan lulusan yang kompeten dan berdaya saing tinggi.")

add_heading_custom(doc, "17.2 Rekomendasi Akhir", level=2)

add_body(doc, "Berdasarkan hasil evaluasi diri, berikut adalah rekomendasi akhir untuk peningkatan mutu Universitas Tulungagung:")

recommendations = [
    "Prioritaskan perbaikan 18 indikator yang berstatus Belum Tercapai, dengan fokus pada Pelayanan Mahasiswa (7 indikator) dan VMTS (5 indikator) yang memiliki jumlah kelemahan terbanyak.",
    "Implementasikan Rencana Tindak Lanjut (RTL) dengan target waktu yang jelas dan penanggung jawab yang ditetapkan, dengan monitoring melalui AMI dan RTM tahunan.",
    "Kembangkan AMI khusus untuk Kerjasama pada tahun 2026, untuk evaluasi yang lebih komprehensif terhadap implementasi MoU/MoA.",
    "Percepat promosi jabatan akademik dosen dan sertifikasi pendidik, dengan target Lektor+ ≥65% dan Serdos ≥96% pada 2026.",
    "Implementasikan One-Stop Service untuk perbaikan pelayanan mahasiswa, dengan target 20/20 indikator mencapai skor ≥3,50 pada 2026.",
    "Diversifikasi sumber pendapatan melalui RGU dan kerja sama DUDI, dengan target opini audit WTP pada 2026.",
    "Kembangkan dashboard SPMI berbasis IT yang terintegrasi dengan PD Dikti untuk monitoring real-time ketercapaian IKU dan standar SPMI.",
    "Mendorong lebih banyak program studi untuk mengajukan akreditasi unggul, dengan target 40% prodi unggul pada 2026 dan 70% pada 2030.",
    "Perkuat kerja sama internasional, khususnya dengan PT di Asia Tenggara, Australia, dan Eropa, untuk meningkatkan reputasi global.",
    "Lakukan review VMTS setiap 5 tahun untuk memastikan relevansi dengan dinamika eksternal, dengan melibatkan DUDI dan asosiasi program studi.",
]

for i, rec in enumerate(recommendations):
    add_numbered(doc, rec)

add_heading_custom(doc, "17.3 Penutup", level=2)

add_body(doc, "Laporan Evaluasi Diri ini merupakan dokumen yang transparan, objektif, dan berbasis bukti (evidence-based), disusun dengan mengacu pada hasil AMI 2024/2025 dan dokumen SPMI Universitas Tulungagung Tahun 2025. Laporan ini mencerminkan komitmen Universitas Tulungagung untuk terus meningkatkan mutu secara berkelanjutan, dalam rangka mewujudkan visi sebagai perguruan tinggi yang berkualitas dan mampu berkompetisi berskala nasional dan internasional.")

add_body(doc, "Universitas Tulungagung berkomitmen untuk melaksanakan seluruh rekomendasi yang tertuang dalam laporan ini, dengan monitoring melalui AMI tahunan dan RTM. Dengan sistem SPMI yang kuat, data AMI yang komprehensif, dan komitmen seluruh sivitas akademika, Universitas Tulungagung siap untuk menghadapi proses akreditasi perguruan tinggi sesuai IAPT 4.1 dan terus meningkatkan mutu penyelenggaraan Tridharma Perguruan Tinggi.")

add_spacer(doc, 3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- END OF REPORT ---")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nLaporan Evaluasi Diri Universitas Tulungagung\nAkreditasi Perguruan Tinggi (APT) 4.1\nTahun Akademik 2024/2025\n\nPusat Penjaminan Mutu (PPM)\nUniversitas Tulungagung\nSK Rektor Nomor: A/002.I/KEP/UNITA/I/2025\nTanggal Penetapan: 1 September 2025")
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

# ============================================================
# SAVE DOCUMENT
# ============================================================

doc.save(output_path)
print(f"Dokumen berhasil dibuat: {output_path}")

# Check file size
import os
size = os.path.getsize(output_path)
print(f"Ukuran file: {size / 1024:.2f} KB")

# ============================================================
# TAMBAHAN: DETAILED INDIKATOR ANALYSIS PER KRITERIA
# ============================================================

# Read the document back and add more content
# Actually, let's add more content before save
# The save already happened, so let's reopen and add

from docx import Document as Doc2
doc2 = Doc2(output_path)

# Add more detailed analysis sections
add_page_break(doc2)

add_heading_custom(doc2, "BAB XVIII\nANALISIS DETAIL INDIKATOR PER KRITERIA", level=1)

add_body(doc2, "Bab ini memberikan analisis narrative detail untuk setiap indikator AMI yang berstatus Belum Tercapai, serta analisis mendalam untuk indikator-indikator kunci yang berstatus Tercapai dengan skor tertinggi. Analisis ini melengkapi tabel data AMI pada bab-bab sebelumnya dengan interpretasi kontekstual dan rekomendasi spesifik.")

# Kriteria 1 Detail
add_heading_custom(doc2, "18.1 Analisis Detail Kriteria 1: VMTS", level=2)

add_heading_custom(doc2, "18.1.1 Indikator Belum Tercapai", level=3)

add_body(doc2, "Indikator 3: VMTS disosialisasikan secara berkala kepada dosen (Skor: 3,42, Target: ≥3,50). Sosialisasi VMTS saat ini dilakukan melalui rapat rutin dan dokumen tertulis. Namun, metode ini dinilai belum cukup efektif oleh sebagian dosen, terutama dosen muda yang lebih terbiasa dengan media digital. Rekomendasi: mengembangkan sosialisasi melalui media sosial kampus, video pendek, infografis, dan newsletter digital yang dikirim secara berkala. Selain itu, perlu dibuat handbook VMTS dalam format yang ringkas dan mudah dibawa.")

add_body(doc2, "Indikator 4: Media dan metode sosialisasi VMTS mudah dipahami (Skor: 3,38, Target: ≥3,50). Skor ini adalah yang terendah di antara semua indikator VMTS. Hal ini menunjukkan bahwa metode sosialisasi yang digunakan saat ini belum optimal dalam menyampaikan substansi VMTS kepada sivitas akademika. Banyak dosen merasa bahwa dokumen VMTS terlalu panjang dan formal, sehingga pesan intinya tidak tertangkap. Rekomendasi: menyederhanakan penyampaian VMTS melalui summary satu halaman, infografis, dan video animasi singkat yang menjelaskan visi, misi, tujuan, dan strategi dalam waktu kurang dari 5 menit.")

add_body(doc2, "Indikator 11: Dosen dilibatkan dalam penyusunan atau evaluasi VMTS (Skor: 3,48, Target: ≥3,50). Keterlibatan dosen dalam penyusunan VMTS masih terbatas pada dosen senior dan pejabat struktural. Dosen muda dan dosen tidak tetap belum terlibat secara optimal. Rekomendasi: membuka forum online untuk input dosen dalam proses review VMTS, serta melibatkan perwakilan dosen muda dalam tim penyusun VMTS.")

add_body(doc2, "Indikator 12: Dosen dilibatkan dalam evaluasi ketercapaian VMTS (Skor: 3,36, Target: ≥3,50). Skor ini adalah yang terendah kedua di antara semua indikator VMTS. Evaluasi ketercapaian VMTS saat ini dilakukan oleh Senat dan PPM, tanpa melibatkan dosen secara luas. Rekomendasi: menyelenggarakan survey tahunan untuk evaluasi ketercapaian VMTS oleh seluruh dosen, dan membahas hasilnya dalam forum terbuka.")

add_body(doc2, "Indikator 14: VMTS relevan dengan kebutuhan dunia kerja (Skor: 3,48, Target: ≥3,50). Beberapa dosen merasa bahwa VMTS belum sepenuhnya mencerminkan kebutuhan DUDI. Rekomendasi: melibatkan DUDI dalam review VMTS melalui focus group discussion, dan menambahkan aspek kewirausahaan dan digital literacy dalam rumusan VMTS.")

add_body(doc2, "Indikator 15: Program universitas mendukung pencapaian VMTS (Skor: 3,40, Target: ≥3,50). Beberapa program kerja unit belum sepenuhnya selaras dengan VMTS. Rekomendasi: melakukan penyelarasan program kerja tahunan dengan VMTS melalui template Renop yang wajib mencantumkan keterkaitan dengan VMTS.")

add_heading_custom(doc2, "18.1.2 Indikator Tercapai dengan Skor Tertinggi", level=3)

add_body(doc2, "Indikator 9: Pimpinan menunjukkan komitmen yang kuat dalam VMTS (Skor: 3,62). Rektor dan pimpinan universitas secara konsisten mengintegrasikan VMTS dalam setiap kebijakan dan program kerja. Komitmen ini ditunjukkan melalui keputusan kebijakan, alokasi anggaran, dan komunikasi rutin dengan sivitas akademika. Ini menjadi kekuatan yang perlu dipertahankan.")

add_body(doc2, "Indikator 5: VMTS menjadi acuan pelaksanaan pendidikan (Skor: 3,62). VMTS telah diintegrasikan dalam kurikulum, RPS, dan proses pembelajaran. Hal ini menunjukkan bahwa VMTS tidak hanya menjadi dokumen formal, tetapi benar-benar menjadi acuan operasional dalam pelaksanaan pendidikan.")

# Kriteria 4 Detail
add_heading_custom(doc2, "18.2 Analisis Detail Kriteria 4: SDM", level=2)

add_heading_custom(doc2, "18.2.1 Indikator Belum Tercapai", level=3)

add_body(doc2, "Indikator 2: Persentase dosen tetap dengan jabatan fungsional minimal Lektor (Capaian: 52,40%, Target: ≥75%). Sebanyak 47,60% dosen tetap masih berada di jabatan Asisten Ahil. Hal ini disebabkan oleh beberapa faktor: (1) banyak dosen muda yang baru bergabung dan belum memenuhi syarat untuk promosi ke Lektor; (2) proses promosi jabatan akademik memerlukan waktu yang lama dan persyaratan yang kompleks; (3) beberapa dosen belum memenuhi syarat publikasi yang dipersyaratkan untuk promosi. Rekomendasi: menyelenggarakan workshop persiapan promosi jabatan akademik secara berkala, menyediakan mentor bagi dosen yang akan promosi, dan memberikan insentif bagi dosen yang berhasil promosi.")

add_body(doc2, "Indikator 3: Kepemilikan sertifikat pendidik atau sertifikat profesi oleh dosen (Capaian: 94,60%, Target: 100%). Sebanyak 5,40% dosen belum memiliki sertifikat pendidik. Faktor penyebab: (1) beberapa dosen baru yang belum sempat mengikuti sertifikasi; (2) proses sertifikasi memerlukan persyaratan yang kompleks dan waktu yang lama; (3) biaya sertifikasi yang tidak murah. Rekomendasi: memfasilitasi dosen yang belum bersertifikat untuk mengikuti pelatihan Serdos, memberikan dukungan finansial untuk biaya sertifikasi, dan menetapkan target tahunan untuk pencapaian 100% sertifikasi.")

add_heading_custom(doc2, "18.2.2 Profil SDM Universitas Tulungagung", level=3)

add_body(doc2, "Berdasarkan data Tahun 2025, Universitas Tulungagung memiliki total 93 orang SDM. Dari jumlah tersebut, seluruh dosen (100%) memiliki kualifikasi minimal S2, dan seluruh tenaga kependidikan (100%) memiliki kualifikasi minimal D3. Sebanyak 78,40% dosen tetap memiliki jabatan fungsional minimal Lektor (meskipun baru 52,40% yang memenuhi target AMI ≥75%). Kepemilikan sertifikat pendidik telah mencapai 94,60%, dan persentase dosen tidak tetap dengan sertifikat kompetensi berada pada 16,20% (di bawah batas maksimum 20%). Realisasi anggaran pengembangan SDM mencapai 88,70% dari total kebutuhan, berada dalam rentang target 85-95%.")

add_body(doc2, "Untuk tenaga kependidikan, seluruh tendik memiliki kualifikasi minimal D3 (100%), keanggotaan tenaga perpustakaan profesional mencapai 100%, dan pelaksanaan uji kompetensi keahlian mencapai 100%. Ketersediaan laporan monev kinerja dosen dan tendik mencapai 100%, demikian juga ketersediaan dokumen DP3. Ketersediaan dokumen Renstra/Renop pengelolaan SDM mencapai 100%, dan ketersediaan Pedoman Kesejahteraan Karyawan juga mencapai 100%.")

add_body(doc2, "Untuk kesejahteraan, AMI Kesejahteraan 2024/2025 menunjukkan rata-rata capaian 96,32% (Sangat Baik) dari 25 indikator, yang seluruhnya berstatus Tercapai. Capaian tertinggi meliputi: ketersediaan Buku Pedoman Kesejahteraan (100%), SOP Pembayaran Gaji (100%), prosedur monitoring (100%), jaminan kesehatan/asuransi (100%), dan ketersediaan instrumen survei (100%). Capaian terendah namun masih Tercapai: fasilitas bantuan anak karyawan (90,80%) dan kecukupan dana pensiun (92,60%).")

# Kriteria 10 Detail
add_heading_custom(doc2, "18.3 Analisis Detail Kriteria 10: Pelayanan Mahasiswa", level=2)

add_heading_custom(doc2, "18.3.1 Indikator Belum Tercapai", level=3)

add_body(doc2, "Indikator 4: Kecepatan respons UPPS terhadap keluhan akademik (Skor: 3,47, Target: ≥3,50). Mahasiswa melaporkan bahwa respons terhadap keluhan akademik kadang memakan waktu lebih lama dari yang diharapkan. Rekomendasi: menetapkan SLA (Service Level Agreement) untuk respons keluhan, maksimal 2 hari kerja, dengan sistem tiket online.")

add_body(doc2, "Indikator 8: Transparansi penggunaan dana/biaya kuliah (Skor: 3,46, Target: ≥3,50). Mahasiswa merasa bahwa informasi tentang penggunaan dana kuliah belum cukup transparan. Rekomendasi: mempublikasikan laporan keuangan ringkas yang dapat dipahami mahasiswa, serta menyelenggarakan forum tahunan untuk menjelaskan alokasi dana.")

add_body(doc2, "Indikator 10: Kecepatan proses verifikasi/pengembalian dana (Skor: 3,48, Target: ≥3,50). Proses verifikasi dan pengembalian dana memerlukan waktu yang dirasa lama oleh mahasiswa. Rekomendasi: menyederhanakan prosedur verifikasi dana dan menetapkan target waktu maksimal 7 hari kerja.")

add_body(doc2, "Indikator 11: Akses informasi kegiatan organisasi mahasiswa (Skor: 3,49, Target: ≥3,50). Informasi tentang kegiatan organisasi mahasiswa dan event kampus belum tersebar merata. Rekomendasi: mengembangkan aplikasi mobile kampus yang mencakup kalender event dan informasi organisasi mahasiswa.")

add_body(doc2, "Indikator 16: Jadwal pembimbingan akademik periodik terstruktur (Skor: 3,48, Target: ≥3,50). Pembimbingan akademik belum sepenuhnya terstruktur dan periodik. Rekomendasi: menetapkan jadwal pembimbingan akademik minimal 2 kali per semester, dengan log book pembimbingan yang terdokumentasi.")

add_body(doc2, "Indikator 18: Pendampingan PKM oleh dosen pembimbing (Skor: 3,49, Target: ≥3,50). Pendampingan PKM oleh dosen pembimbing belum optimal. Rekomendasi: menetapkan rasio dosen pembimbing PKM yang wajar, dan memberikan insentif bagi dosen yang aktif membimbing PKM.")

add_body(doc2, "Indikator 19: Fasilitas kebersihan (toilet, tempat sampah) (Skor: 3,44, Target: ≥3,50). Fasilitas kebersihan mendapat skor terendah di antara semua indikator pelayanan mahasiswa. Hal ini menunjukkan bahwa kebersihan toilet dan ketersediaan tempat sampah masih menjadi masalah. Rekomendasi: meningkatkan frekuensi pembersihan toilet, menambah jumlah tempat sampah, dan melibatkan mahasiswa dalam program kebersihan kampus.")

add_heading_custom(doc2, "18.3.2 Tren Peningkatan Pelayanan Mahasiswa", level=3)

add_body(doc2, "Meskipun terdapat 7 indikator yang belum tercapai, penting untuk dicatat bahwa terdapat tren peningkatan yang signifikan. Pada Ganjil 2023-1, hanya 8 dari 20 indikator yang mencapai skor ≥3,50. Pada Ganjil 2024-1, jumlah ini meningkat menjadi 11 indikator. Pada Genap 2024-2, jumlah indikator yang mencapai target melonjak menjadi 17 dari 20. Skor rata-rata juga meningkat dari 3,42 menjadi 3,48 menjadi 3,58. Peningkatan ini merupakan bukti efektivitas implementasi Rencana Tindak Lanjut (RTL) dari periode sebelumnya.")

add_body(doc2, "Indikator dengan skor tertinggi pada Genap 2024-2: akses SIAKAD (3,70), informasi jadwal perkuliahan (3,68), dan kemudahan akses dokumen akademik (3,56). Hal ini menunjukkan bahwa sistem informasi akademik digital berjalan dengan baik. Indikator dengan peningkatan terbesar: layanan beasiswa (+0,05 dari Ganjil ke Genap), akses perpustakaan, dan akses laboratorium.")

# Kriteria 6 Detail - Pendidikan
add_heading_custom(doc2, "18.4 Analisis Detail Kriteria 6: Pendidikan", level=2)

add_heading_custom(doc2, "18.4.1 Analisis Capaian Lulusan", level=3)

add_body(doc2, "Capaian lulusan Universitas Tulungagung menunjukkan hasil yang sangat baik. Persentase lulusan terserap di dunia kerja mencapai 86,50%, jauh di atas target 70%. Ini berarti dari setiap 100 lulusan, sekitar 87 orang mendapatkan pekerjaan dalam waktu kurang dari 6 bulan setelah kelulusan. Capaian ini meningkat dari 80% pada tahun 2023, menunjukkan tren positif.")

add_body(doc2, "Persentase lulusan yang menjadi wirausahawan mencapai 18,20%, di atas target 10%. Ini berarti dari setiap 100 lulusan, sekitar 18 orang memilih untuk memulai bisnis sendiri. Capaian ini meningkat dari 15% pada tahun 2023, menunjukkan bahwa program kewirausahaan di Universitas Tulungagung mulai membuahkan hasil.")

add_body(doc2, "Rata-rata IPK lulusan adalah 3,48, di atas target 3,00. Ini menunjukkan bahwa secara akademik, lulusan Universitas Tulungagung memiliki kompetensi yang baik. Rata-rata masa studi 3,8 tahun (di bawah target 4,5 tahun) menunjukkan bahwa sebagian besar mahasiswa dapat menyelesaikan studi tepat waktu. Rata-rata masa tunggu kerja 3,2 bulan (jauh di bawah target 6 bulan) menunjukkan bahwa lulusan Universitas Tulungagung memiliki daya saing yang tinggi di pasar kerja.")

add_heading_custom(doc2, "18.4.2 Analisis Proses Pembelajaran", level=3)

add_body(doc2, "AMI Proses Pembelajaran dilaksanakan dua kali per tahun (Ganjil dan Genap), dengan 20 indikator yang dievaluasi menggunakan skala Likert 1-4. Pada Genap 2024-2, seluruh 20 indikator berstatus Tercapai (skor ≥3,50). Skor tertinggi adalah penguasaan materi perkuliahan oleh dosen (3,82), yang menunjukkan bahwa dosen Universitas Tulungagung memiliki kompetensi yang sangat baik dalam bidang keilmuannya.")

add_body(doc2, "Skor terendah (namun masih Tercapai) adalah efektivitas metode pembelajaran daring (3,65). Meskipun skor ini berada di atas target, metode daring masih menjadi area yang perlu peningkatan dibandingkan metode luring (3,76) dan hybrid (3,74). Rekomendasi: menyelenggarakan pelatihan khusus untuk metode pembelajaran daring, mengembangkan konten digital yang lebih interaktif, dan memberikan dukungan teknis bagi dosen dalam penggunaan LMS.")

add_body(doc2, "Tren peningkatan skor dari 3,43 (Ganjil 2023-2) ke 3,49 (Ganjil 2024-1) ke 3,59 (Genap 2024-2) menunjukkan efektivitas implementasi RTL. Seluruh fakultas (FE, FH, FISIP, FP, FT, F.KES) menunjukkan peningkatan performa. Praktik baik dari Prodi Akuntansi (skor tertinggi) dapat direplikasi ke Prodi Manajemen dan prodi lainnya.")

add_heading_custom(doc2, "18.4.3 Analisis Kurikulum dan Isi", level=3)

add_body(doc2, "Ketersediaan dan kelengkapan dokumen kurikulum mencapai 100%, dengan frekuensi tinjauan kurikulum 2 kali per tahun. Rasio mahasiswa per dosen pembimbing akademik adalah 16:1 (target ≤20:1), dan rasio mahasiswa per dosen pembimbing tugas akhir adalah 8:1 (target ≤10:1). Peningkatan rasio pembimbing tugas akhir dari 12:1 menjadi 8:1 menunjukkan komitmen Universitas Tulungagung dalam meningkatkan kualitas pembimbingan tugas akhir.")

add_body(doc2, "Persentase mata kuliah blended learning mencapai 32% (target ≥20%), menunjukkan bahwa Universitas Tulungagung telah mengadopsi pembelajaran hybrid. Namun, masih ada ruang untuk peningkatan, dengan target 40% pada 2026. Rata-rata jumlah mahasiswa per kelas adalah 32 (target ≤40), yang menunjukkan ukuran kelas yang optimal untuk pembelajaran interaktif.")

add_page_break(doc2)

# GLOSSARY
add_heading_custom(doc2, "GLOSARIUM", level=1)

add_body(doc2, "Berikut adalah daftar istilah dan singkatan yang digunakan dalam Laporan Evaluasi Diri ini:", indent=False)

glossary = [
    ("AMI", "Audit Mutu Internal — kegiatan pemeriksaan kepatuhan yang dilakukan oleh auditor internal universitas terhadap pelaksanaan SPMI"),
    ("AEE PT", "Angka Efisiensi Edukasi Perguruan Tinggi — persentase mahasiswa yang lulus tepat waktu"),
    ("APT", "Akreditasi Perguruan Tinggi — proses penilaian kelayakan program studi dan perguruan tinggi"),
    ("BAN-PT", "Badan Akreditasi Nasional Perguruan Tinggi — badan yang dibentuk pemerintah untuk mengembangkan sistem akreditasi"),
    ("BEM", "Badan Eksekutif Mahasiswa — organisasi eksekutif mahasiswa tingkat universitas"),
    ("CPL", "Capaian Pembelajaran Lulusan — rumusan kompetensi lulusan"),
    ("CPMK", "Capaian Pembelajaran Mata Kuliah — rumusan kompetensi mata kuliah"),
    ("DUDI", "Dunia Usaha, Dunia Industri, dan Dunia Kerja"),
    ("GKM", "Gugus Penjaminan Mutu — unsur di tingkat program studi yang mengelola SPMI"),
    ("HIMA", "Himpunan Mahasiswa — organisasi mahasiswa tingkat program studi"),
    ("HKI", "Hak Kekayaan Intelektual"),
    ("IAPT", "Instrumen Akreditasi Perguruan Tinggi"),
    ("IKU", "Indikator Kinerja Utama"),
    ("KKNI", "Kerangka Kualifikasi Nasional Indonesia"),
    ("LAM", "Lembaga Akreditasi Mandiri"),
    ("LED", "Laporan Evaluasi Diri"),
    ("LMS", "Learning Management System"),
    ("LPPM", "Lembaga Penelitian dan Pengabdian kepada Masyarakat"),
    ("LPM", "Lembaga Penjaminan Mutu — lembaga di tingkat fakultas"),
    ("MBKM", "Merdeka Belajar Kampus Merdeka"),
    ("MoU/MoA", "Memorandum of Understanding / Memorandum of Agreement"),
    ("OBE", "Outcome-Based Education"),
    ("PD Dikti", "Pangkalan Data Pendidikan Tinggi"),
    ("PkM", "Pengabdian kepada Masyarakat"),
    ("PPM", "Pusat Penjaminan Mutu"),
    ("PPEPP", "Penetapan, Pelaksanaan, Evaluasi, Pengendalian, Peningkatan"),
    ("PTK", "Permintaan Tindakan Koreksi"),
    ("RENOP", "Rencana Operasional"),
    ("RENSTRA", "Rencana Strategis"),
    ("RIP", "Rencana Induk Pengembangan"),
    ("RPS", "Rencana Pembelajaran Semester"),
    ("RTM", "Rapat Tinjauan Manajemen"),
    ("RTL", "Rencana Tindak Lanjut"),
    ("SDGs", "Sustainable Development Goals"),
    ("SIAKAD", "Sistem Informasi Akademik"),
    ("SKPI", "Surat Keterangan Pendamping Ijazah"),
    ("SN Dikti", "Standar Nasional Pendidikan Tinggi"),
    ("SOP", "Standar Operasional Prosedur"),
    ("SPI", "Satuan Pengawas Internal"),
    ("SPM Dikti", "Sistem Penjaminan Mutu Pendidikan Tinggi"),
    ("SPMI", "Sistem Penjaminan Mutu Internal"),
    ("SPME", "Sistem Penjaminan Mutu Eksternal"),
    ("SOTK", "Struktur Organisasi dan Tata Kerja"),
    ("TTG", "Teknologi Tepat Guna"),
    ("UKM", "Unit Kegiatan Mahasiswa"),
    ("UPPS", "Unit Pengelola Program Studi"),
    ("VMTS", "Visi, Misi, Tujuan, dan Strategi"),
]

for term, definition in glossary:
    p = doc2.add_paragraph()
    run1 = p.add_run(f"{term}")
    run1.font.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = NAVY
    run1.font.name = 'Calibri'
    run2 = p.add_run(f" — {definition}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK_GRAY
    run2.font.name = 'Calibri'
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.space_after = Pt(3)

add_page_break(doc2)

# DAFTAR SINGKATAN
add_heading_custom(doc2, "DAFTAR SINGKATAN", level=1)

add_body(doc2, "Daftar singkatan yang digunakan dalam laporan ini telah dicantumkan dalam Glosarium. Singkatan-singkatan utama yang sering muncul: AMI (Audit Mutu Internal), APT (Akreditasi Perguruan Tinggi), BAN-PT (Badan Akreditasi Nasional Perguruan Tinggi), DUDI (Dunia Usaha, Dunia Industri, dan Dunia Kerja), GKM (Gugus Penjaminan Mutu), IAPT (Instrumen Akreditasi Perguruan Tinggi), IKU (Indikator Kinerja Utama), LAM (Lembaga Akreditasi Mandiri), LED (Laporan Evaluasi Diri), LPPM (Lembaga Penelitian dan Pengabdian kepada Masyarakat), LPM (Lembaga Penjaminan Mutu), MBKM (Merdeka Belajar Kampus Merdeka), PD Dikti (Pangkalan Data Pendidikan Tinggi), PkM (Pengabdian kepada Masyarakat), PPM (Pusat Penjaminan Mutu), PPEPP (Penetapan, Pelaksanaan, Evaluasi, Pengendalian, Peningkatan), RENOP (Rencana Operasional), RENSTRA (Rencana Strategis), RPS (Rencana Pembelajaran Semester), RTM (Rapat Tinjauan Manajemen), RTL (Rencana Tindak Lanjut), SIAKAD (Sistem Informasi Akademik), SN Dikti (Standar Nasional Pendidikan Tinggi), SPM Dikti (Sistem Penjaminan Mutu Pendidikan Tinggi), SPMI (Sistem Penjaminan Mutu Internal), SPME (Sistem Penjaminan Mutu Eksternal), UPPS (Unit Pengelola Program Studi), VMTS (Visi, Misi, Tujuan, dan Strategi).")

add_spacer(doc2, 3)

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- END OF REPORT ---")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.italic = True
run.font.name = 'Calibri'

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\nLaporan Evaluasi Diri Universitas Tulungagung\nAkreditasi Perguruan Tinggi (APT) 4.1\nTahun Akademik 2024/2025\n\nPusat Penjaminan Mutu (PPM)\nUniversitas Tulungagung\nSK Rektor Nomor: A/002.I/KEP/UNITA/I/2025\nTanggal Penetapan: 1 September 2025")
run.font.size = Pt(10)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

doc2.save(output_path)
print(f"Dokumen updated: {output_path}")
size = os.path.getsize(output_path)
print(f"Ukuran file: {size / 1024:.2f} KB")
