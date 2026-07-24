"""
Generate LAPORAN HASIL PELAKSANAAN DAN EVALUASI ASSESSMENT AMI 2024-2025
Dokumen untuk tahap Evaluasi (E) siklus PPEPP SPMI
Sumber: 15 dokumen AMI 2024-2025, 33 Standar SPMI, Permen 39/2025
"""
import json, os
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Load data
with open('/tmp/standar_2025.json') as f:
    STANDAR = json.load(f)
with open('/tmp/dokumen_ami_2025.json') as f:
    DOKUMEN = json.load(f)

# Calculate summary
RATA_SKOR = sum(s['skor'] for s in STANDAR) / len(STANDAR)
STATUS_COUNTS = {}
for s in STANDAR:
    STATUS_COUNTS[s['status']] = STATUS_COUNTS.get(s['status'], 0) + 1
TOTAL_TEMUAN = sum(s['temuan'] for s in STANDAR)
TOTAL_PTK = sum(s['ptk'] for s in STANDAR)
RATA_TL = sum(s['tindakLanjut'] for s in STANDAR) / len(STANDAR)

# Colors
NAVY = RGBColor(0x1A, 0x3A, 0x6E)
NAVY_LIGHT = RGBColor(0x38, 0x6E, 0xE6)
EMERALD = RGBColor(0x10, 0x98, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
ROSE = RGBColor(0xF4, 0x3F, 0x5E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1F, 0x29, 0x37)

# Status colors
STATUS_COLOR = {
    'MS': EMERALD, 'ML': RGBColor(0x06, 0xB6, 0xD4),
    'BS': AMBER, 'MSV': ROSE
}
STATUS_LABEL = {
    'MS': 'Mencapai Standar', 'ML': 'Melampaui Standar',
    'BS': 'Belum mencapai Standar', 'MSV': 'Menyimpang dari Standar'
}

doc = Document()

# ============ PAGE SETUP ============
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

# Helper functions
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="386EE6", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def add_heading(text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    return p

def add_body(text, justify=True, indent=True, bold=False, color=BLACK, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_table_header(table, headers, bg_color="1A3A6E"):
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_shading(cell, bg_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_table_row(table, values, align='left', bold=False, bg=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(str(val))
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        run.bold = bold
        if bg:
            set_cell_shading(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

# ============ COVER PAGE ============
# Spacer
for _ in range(3):
    doc.add_paragraph()

# Logo placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSITAS TULUNGAGUNG')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PUSAT PENJAMINAN MUTU (PPM)')
run.font.size = Pt(13)
run.font.color.rgb = NAVY_LIGHT
run.font.name = 'Calibri'

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(20)
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.color.rgb = NAVY_LIGHT
run.font.size = Pt(10)

# Title
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LAPORAN HASIL PELAKSANAAN DAN\nEVALUASI ASSESSMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('AUDIT MUTU INTERNAL (AMI)')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = NAVY_LIGHT
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TAHUN AKADEMIK 2024/2025')
run.font.size = Pt(14)
run.font.color.rgb = BLACK
run.font.name = 'Calibri'

# Divider
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('Siklus PPEPP — Tahap Evaluasi (E)')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Berbasis Peraturan Menteri Pendidikan Tinggi No. 39 Tahun 2025\ntentang Penjaminan Mutu Pendidikan Tinggi')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# Bottom info
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TULUNGAGUNG\n2025')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

doc.add_page_break()

# ============ LEMBAR PENGESAHAN ============
add_heading('LEMBAR PENGESAHAN', level=1, color=NAVY, size=18)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body('Laporan Hasil Pelaksanaan dan Evaluasi Assessment Audit Mutu Internal (AMI) Tahun Akademik 2024/2025 ini telah disusun berdasarkan hasil audit terhadap 33 Standar SPMI yang mencakup 15 dokumen AMI dengan total 212 indikator penilaian, sebagai bentuk pelaksanaan tahap Evaluasi (E) pada siklus PPEPP Sistem Penjaminan Mutu Internal Universitas Tulungagung.', justify=True, indent=True)

doc.add_paragraph()
doc.add_paragraph()

# Tanda tangan table
table = doc.add_table(rows=6, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
table.autofit = False

# Header
headers = ['Proses', 'Nama Jabatan', 'Nama']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    set_cell_shading(cell, "1A3A6E")

# Data
data = [
    ('Perumus', 'Anggota Audit & Monev-in', 'Desi Rahmawati, SE., M.M'),
    ('Pemeriksa', 'Kabag. Audit & Monev-in', 'Eni Minarni, SE.Ak., M.Ak'),
    ('Persetujuan', 'Ketua Pusat Penjaminan Mutu', 'Dr. Anang Sugeng Cahoyono, S.AP., S.H., M.Si.'),
    ('Penetapan', 'Rektor', 'Dr. Muharsono, M.Si.'),
    ('Pengendalian', 'Wakil Rektor I', 'Dr. Ir. Bambang Tri Kurnianto, Dipl.Inv., M.MA'),
]
for i, (proses, jabatan, nama) in enumerate(data, 1):
    row = table.rows[i]
    for j, val in enumerate([proses, jabatan, nama]):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if j == 0:
            run.bold = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tulungagung, Agustus 2025')
run.font.size = Pt(11)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Pusat Penjaminan Mutu')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'

doc.add_page_break()

# ============ DAFTAR ISI ============
add_heading('DAFTAR ISI', level=1, color=NAVY, size=18)

toc_items = [
    ('LEMBAR PENGESAHAN', 'ii'),
    ('DAFTAR ISI', 'iii'),
    ('DAFTAR TABEL', 'iv'),
    ('BAB I PENDAHULUAN', '1'),
    ('  1.1 Latar Belakang', '1'),
    ('  1.2 Dasar Hukum', '1'),
    ('  1.3 Tujuan Assessment', '2'),
    ('  1.4 Ruang Lingkup Assessment', '2'),
    ('BAB II METODOLOGI ASSESSMENT', '3'),
    ('  2.1 Siklus PPEPP', '3'),
    ('  2.2 Teknik Pengumpulan Data', '3'),
    ('  2.3 Skala Penilaian dan Kategori Status', '4'),
    ('BAB III HASIL ASSESSMENT 33 STANDAR SPMI', '5'),
    ('  3.1 Ringkasan Eksekutif', '5'),
    ('  3.2 Standar Pendidikan (8 Standar)', '6'),
    ('  3.3 Standar Penelitian (8 Standar)', '8'),
    ('  3.4 Standar Pengabdian kepada Masyarakat (8 Standar)', '10'),
    ('  3.5 Standar Tambahan (9 Standar)', '12'),
    ('BAB IV REKAPITULASI 15 DOKUMEN AMI', '14'),
    ('BAB V ANALISIS TEMUAN DAN PTK', '16'),
    ('BAB VI KESIMPULAN DAN REKOMENDASI', '18'),
    ('LAMPIRAN: Daftar Dokumen AMI 2024-2025', '20'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    # Tab stop for page number
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    is_bab = item.startswith('BAB') or item.startswith('LAMPIRAN') or item.startswith('LEMBAR') or item.startswith('DAFTAR')
    run = p.add_run(item)
    run.font.size = Pt(11 if is_bab else 10)
    run.bold = is_bab
    run.font.name = 'Calibri'
    run.font.color.rgb = NAVY if is_bab else BLACK
    
    run2 = p.add_run('\t' + page)
    run2.font.size = Pt(11 if is_bab else 10)
    run2.bold = is_bab
    run2.font.name = 'Calibri'
    run2.font.color.rgb = NAVY if is_bab else BLACK

doc.add_page_break()

# ============ DAFTAR TABEL ============
add_heading('DAFTAR TABEL', level=1, color=NAVY, size=18)

tabel_items = [
    ('Tabel 1.1 Dasar Hukum Assessment AMI 2024-2025', '2'),
    ('Tabel 2.1 Skala Penilaian dan Kategori Status AMI', '4'),
    ('Tabel 3.1 Ringkasan Eksekutif Hasil AMI 33 Standar SPMI', '5'),
    ('Tabel 3.2 Hasil Assessment 8 Standar Pendidikan', '6'),
    ('Tabel 3.3 Hasil Assessment 8 Standar Penelitian', '8'),
    ('Tabel 3.4 Hasil Assessment 8 Standar Pengabdian kepada Masyarakat', '10'),
    ('Tabel 3.5 Hasil Assessment 9 Standar Tambahan', '12'),
    ('Tabel 4.1 Rekapitulasi 15 Dokumen AMI 2024-2025', '14'),
    ('Tabel 5.1 Analisis Temuan dan PTK per Kategori Standar', '16'),
    ('Tabel 6.1 Rekomendasi Tindak Lanjut Prioritas', '18'),
]
for item, page in tabel_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLACK
    run2 = p.add_run('\t' + page)
    run2.font.size = Pt(10.5)
    run2.font.name = 'Calibri'

doc.add_page_break()

# ============ BAB I PENDAHULUAN ============
add_heading('BAB I\nPENDAHULUAN', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('1.1 Latar Belakang', level=2)
add_body('Penjaminan mutu pendidikan tinggi merupakan kegiatan sistemik untuk meningkatkan mutu pendidikan tinggi secara berencana dan berkelanjutan. Sistem Penjaminan Mutu Internal (SPMI) di Universitas Tulungagung dilaksanakan melalui siklus PPEPP (Penetapan, Pelaksanaan, Evaluasi, Pengendalian, dan Peningkatan) sebagaimana diamanatkan dalam Peraturan Menteri Pendidikan Tinggi Nomor 39 Tahun 2025 tentang Penjaminan Mutu Pendidikan Tinggi.')

add_body('Tahap Evaluasi (E) dalam siklus PPEPP dilaksanakan melalui Audit Mutu Internal (AMI), yaitu kegiatan audit yang dilakukan oleh unit penjaminan mutu internal untuk menilai tingkat pencapaian standar yang telah ditetapkan. AMI Tahun Akademik 2024/2025 merupakan pelaksanaan audit terhadap 33 Standar SPMI yang mencakup 8 Standar Pendidikan, 8 Standar Penelitian, 8 Standar Pengabdian kepada Masyarakat, dan 9 Standar Tambahan.')

add_body('Laporan Hasil Pelaksanaan dan Evaluasi Assessment ini disusun sebagai dokumen pertanggungjawaban pelaksanaan AMI TA 2024/2025 yang mencakup 15 dokumen AMI dengan total 212 indikator penilaian. Laporan ini menjadi dasar bagi pimpinan universitas untuk mengambil kebijakan strategis dalam tahap Pengendalian (P) dan Peningkatan (P) pada siklus PPEPP berikutnya.')

add_heading('1.2 Dasar Hukum', level=2)
add_body('Pelaksanaan Assessment AMI TA 2024/2025 berlandaskan pada peraturan-peraturan berikut:')

# Tabel dasar hukum
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['No', 'Peraturan'])
hukum = [
    ('1', 'Undang-Undang Republik Indonesia Nomor 12 Tahun 2012 tentang Pendidikan Tinggi'),
    ('2', 'Peraturan Pemerintah Republik Indonesia No. 4 Tahun 2014 tentang Penyelenggaraan Pendidikan Tinggi dan Pengelolaan Perguruan Tinggi'),
    ('3', 'Peraturan Menteri Pendidikan Tinggi Nomor 39 Tahun 2025 tentang Penjaminan Mutu Pendidikan Tinggi'),
    ('4', 'SK Rektor UNITA Nomor A/039.A/KEP/UNITA/VIII/2020 tentang Manual Mutu Monev-In Tri Dharma Perguruan Tinggi'),
    ('5', 'SK Rektor UNITA Nomor A/026/KEP/UNITA/VI/2022 tentang Pengangkatan Pejabat Penjaminan Mutu Universitas Tulungagung'),
]
for no, rule in hukum:
    add_table_row(table, [no, rule], align='left')

doc.add_paragraph()

add_heading('1.3 Tujuan Assessment', level=2)
add_body('Assessment AMI TA 2024/2025 bertujuan untuk:')
add_bullet('Menilai tingkat pencapaian 33 Standar SPMI di Universitas Tulungagung berdasarkan 212 indikator yang terdistribusi dalam 15 dokumen AMI.')
add_bullet('Mengidentifikasi temuan dan Permintaan Tindakan Koreksi (PTK) yang perlu ditindaklanjuti oleh unit pengelola standar.')
add_bullet('Menyediakan informasi yang menjadi dasar pengambilan kebijakan oleh pimpinan universitas dalam tahap Pengendalian dan Peningkatan siklus PPEPP.')
add_bullet('Memenuhi amanat Permen 39/2025 tentang pelaksanaan tahap Evaluasi dalam SPMI.')

add_heading('1.4 Ruang Lingkup Assessment', level=2)
add_body('Ruang lingkup Assessment AMI TA 2024/2025 mencakup:')
add_bullet('33 Standar SPMI: 8 Standar Pendidikan, 8 Standar Penelitian, 8 Standar Pengabdian kepada Masyarakat (PkM), dan 9 Standar Tambahan.')
add_bullet('15 Dokumen AMI: AMI Pendidikan, AMI Proses Pembelajaran Ganjil & Genap, AMI SDM, AMI Kesejahteraan, AMI Keuangan, AMI Sarana dan Prasarana, AMI VMTS, AMI TPTK, AMI Mahasiswa, AMI Penelitian, AMI Pengabdian, AMI MBKM, AMI Pelayanan Mahasiswa Ganjil & Genap.')
add_bullet('212 Indikator penilaian dengan skala pencapaian dan kategori status MS/ML/BS/MSV.')
add_bullet('Periode audit: sepanjang Tahun Akademik 2024/2025 oleh PPM UNITA.')

doc.add_page_break()

# ============ BAB II METODOLOGI ============
add_heading('BAB II\nMETODOLOGI ASSESSMENT', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('2.1 Siklus PPEPP', level=2)
add_body('Sistem Penjaminan Mutu Internal (SPMI) Universitas Tulungagung dilaksanakan melalui siklus PPEPP (Penetapan, Pelaksanaan, Evaluasi, Pengendalian, dan Peningkatan) sebagaimana diatur dalam Permen 39/2025. Assessment AMI TA 2024/2025 ini merupakan pelaksanaan tahap Evaluasi (E) yang berfungsi untuk:')

add_bullet('Penetapan (P): Penetapan 33 Standar SPMI melalui 4 Dokumen Mutu (Kebijakan, Manual, Standar, Formulir) yang ditetapkan 1 September 2025.')
add_bullet('Pelaksanaan (P): Implementasi 33 standar di seluruh unit kerja universitas sepanjang TA 2024/2025.')
add_bullet('Evaluasi (E): Audit Mutu Internal melalui 15 dokumen AMI dengan 212 indikator — INI TAHAP YANG DILAPORKAN.')
add_bullet('Pengendalian (P): Rapat Tinjauan Manajemen (RTM) untuk merumuskan Rencana Tindak Lanjut (RTL).')
add_bullet('Peningkatan (P): Implementasi RTL untuk peningkatan mutu berkelanjutan.')

add_heading('2.2 Teknik Pengumpulan Data', level=2)
add_body('Pengumpulan data AMI TA 2024/2025 dilakukan melalui kombinasi pendekatan kuantitatif dan kualitatif dengan instrumen:')

add_bullet('Kuesioner online (Google Form) dengan skala Likert 1-4 untuk standar berbasis persepsi (VMTS, TPTK, Mahasiswa, Pelayanan, Kesejahteraan, Proses Pembelajaran).')
add_bullet('Checklist dokumen untuk standar berbasis dokumen (Pendidikan, Penelitian, PkM, SDM, Keuangan, Sarpras, MBKM).')
add_bullet('Wawancara terstruktur dengan penanggung jawab standar di setiap unit kerja.')
add_bullet('Observasi lapangan untuk sarana dan prasarana.')
add_bullet('Triangulasi data dari PD-Dikti, evaluasi internal, dan stakeholder eksternal.')

add_heading('2.3 Skala Penilaian dan Kategori Status', level=2)
add_body('Penilaian AMI menggunakan skala pencapaian 0-100 dengan 4 kategori status sebagaimana ditunjukkan pada Tabel 2.1.')

# Tabel skala
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['Kategori Status', 'Rentang Skor', 'Keterangan'])
skala = [
    ('MS (Mencapai Standar)', '75 - 89', 'Capaian memenuhi standar yang ditetapkan'),
    ('ML (Melampaui Standar)', '90 - 100', 'Capaian melebihi standar yang ditetapkan'),
    ('BS (Belum mencapai Standar)', '60 - 74', 'Capaian belum memenuhi standar, perlu PTK'),
    ('MSV (Menyimpang dari Standar)', '< 60', 'Capaian menyimpang jauh, perlu tindakan korektif segera'),
]
for kat, rentang, ket in skala:
    row = add_table_row(table, [kat, rentang, ket], align='left')
    # Color the status cell
    status_key = kat.split(' ')[0]
    if status_key in STATUS_COLOR:
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = STATUS_COLOR[status_key]
        row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ============ BAB III HASIL ASSESSMENT ============
add_heading('BAB III\nHASIL ASSESSMENT 33 STANDAR SPMI', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('3.1 Ringkasan Eksekutif', level=2)
add_body(f'Hasil Assessment AMI TA 2024/2025 terhadap 33 Standar SPMI Universitas Tulungagung menunjukkan capaian yang sangat baik. Rata-rata skor AMI institusi mencapai {RATA_SKOR:.1f}/100 dengan distribusi status sebagai berikut: {STATUS_COUNTS.get("MS", 0)} standar Mencapai Standar (MS), {STATUS_COUNTS.get("ML", 0)} standar Melampaui Standar (ML), {STATUS_COUNTS.get("BS", 0)} standar Belum mencapai Standar (BS), dan {STATUS_COUNTS.get("MSV", 0)} standar Menyimpang dari Standar (MSV).')

add_body(f'Total temuan yang teridentifikasi sebanyak {TOTAL_TEMUAN} temuan dengan {TOTAL_PTK} Permintaan Tindakan Koreksi (PTK). Rata-rata tindak lanjut mencapai {RATA_TL:.1f}%, menunjukkan komitmen yang tinggi dari seluruh unit pengelola standar dalam menindaklanjuti hasil audit.')

# Tabel ringkasan eksekutif
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['Indikator', 'Capaian'])
ringkasan = [
    ('Total Standar SPMI Diaudit', '33 standar'),
    ('Total Dokumen AMI', '15 dokumen'),
    ('Total Indikator Penilaian', '212 indikator'),
    ('Rata-rata Skor AMI Institusi', f'{RATA_SKOR:.1f}/100'),
    ('Standar MS (Mencapai Standar)', f'{STATUS_COUNTS.get("MS", 0)} standar'),
    ('Standar ML (Melampaui Standar)', f'{STATUS_COUNTS.get("ML", 0)} standar'),
    ('Standar BS (Belum mencapai Standar)', f'{STATUS_COUNTS.get("BS", 0)} standar'),
]
for ind, cap in ringkasan:
    add_table_row(table, [ind, cap], align='left')

doc.add_paragraph()

# Function to create standar table
def create_standar_table(standar_list, title):
    add_heading(title, level=2)
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['No', 'Nama Standar', 'Kode', 'Skor', 'Status', 'Temuan', 'PTK', 'TL (%)'])
    
    for s in standar_list:
        row = table.add_row()
        values = [s['no'], s['nama'], s['kode'], s['skor'], s['status'], s['temuan'], s['ptk'], f"{s['tindakLanjut']}%"]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            # Color skor
            if i == 3:
                if s['skor'] >= 90:
                    run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)
                elif s['skor'] >= 75:
                    run.font.color.rgb = EMERALD
                else:
                    run.font.color.rgb = AMBER
                run.bold = True
            # Color status
            if i == 4:
                run.font.color.rgb = STATUS_COLOR.get(s['status'], BLACK)
                run.bold = True
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_paragraph()
    
    # Catatan detail
    add_body('Catatan Hasil Assessment:', bold=True, indent=False)
    for s in standar_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{s['no']}. {s['nama']} ({s['kode']}): ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = NAVY
        run2 = p.add_run(f"Skor {s['skor']} ({STATUS_LABEL[s['status']]}). {s['catatan']}")
        run2.font.size = Pt(10)
        run2.font.name = 'Calibri'
    
    doc.add_paragraph()

# 3.2 Standar Pendidikan
pendidikan = [s for s in STANDAR if s['kategori'] == 'Pendidikan']
create_standar_table(pendidikan, '3.2 Standar Pendidikan (8 Standar)')

doc.add_page_break()

# 3.3 Standar Penelitian
penelitian = [s for s in STANDAR if s['kategori'] == 'Penelitian']
create_standar_table(penelitian, '3.3 Standar Penelitian (8 Standar)')

doc.add_page_break()

# 3.4 Standar PkM
pkm = [s for s in STANDAR if s['kategori'] == 'PkM']
create_standar_table(pkm, '3.4 Standar Pengabdian kepada Masyarakat (8 Standar)')

doc.add_page_break()

# 3.5 Standar Tambahan
tambahan = [s for s in STANDAR if s['kategori'] == 'Tambahan']
create_standar_table(tambahan, '3.5 Standar Tambahan (9 Standar)')

doc.add_page_break()

# ============ BAB IV REKAPITULASI DOKUMEN AMI ============
add_heading('BAB IV\nREKAPITULASI 15 DOKUMEN AMI', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body('Assessment AMI TA 2024/2025 dilaksanakan melalui 15 dokumen AMI yang mencakup seluruh aspek Tridharma Perguruan Tinggi dan standar tambahan. Berikut rekapitulasi 15 dokumen AMI tersebut:')

# Tabel dokumen
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['No', 'Nama Dokumen AMI', 'Kode Dokumen', 'Tanggal Audit', 'Kategori'])

for i, d in enumerate(DOKUMEN, 1):
    add_table_row(table, [i, d['nama'], d['kode'], d['tanggal'], d['kategori']], align='left')

doc.add_paragraph()

add_body('Ke-15 dokumen AMI tersebut mencakup total 212 indikator penilaian yang terdistribusi pada 33 Standar SPMI. Setiap dokumen AMI diaudit oleh Tim Audit Mutu Internal PPM UNITA dengan melibatkan Gugus Penjaminan Mutu (GPM) di tingkat fakultas dan program studi.')

doc.add_page_break()

# ============ BAB V ANALISIS TEMUAN ============
add_heading('BAB V\nANALISIS TEMUAN DAN PTK', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('5.1 Analisis Temuan per Kategori Standar', level=2)
add_body(f'Hasil Assessment AMI TA 2024/2025 mengidentifikasi total {TOTAL_TEMUAN} temuan dengan {TOTAL_PTK} Permintaan Tindakan Koreksi (PTK) yang tersebar pada 33 Standar SPMI. Distribusi temuan dan PTK per kategori standar disajikan pada Tabel 5.1.')

# Tabel analisis per kategori
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['Kategori', 'Jumlah Standar', 'Rata-rata Skor', 'Total Temuan', 'Total PTK'])

kategori_stats = {}
for s in STANDAR:
    if s['kategori'] not in kategori_stats:
        kategori_stats[s['kategori']] = {'count': 0, 'skor': 0, 'temuan': 0, 'ptk': 0}
    kategori_stats[s['kategori']]['count'] += 1
    kategori_stats[s['kategori']]['skor'] += s['skor']
    kategori_stats[s['kategori']]['temuan'] += s['temuan']
    kategori_stats[s['kategori']]['ptk'] += s['ptk']

for kat, stats in kategori_stats.items():
    rata = stats['skor'] / stats['count']
    add_table_row(table, [kat, stats['count'], f"{rata:.1f}", stats['temuan'], stats['ptk']], align='center')

# Total row
total_rata = RATA_SKOR
add_table_row(table, ['TOTAL', '33', f"{total_rata:.1f}", TOTAL_TEMUAN, TOTAL_PTK], align='center', bold=True, bg="E8F0FE")

doc.add_paragraph()

add_heading('5.2 Identifikasi Standar dengan PTK Tertinggi', level=2)
add_body('Standar dengan jumlah PTK tertinggi memerlukan perhatian khusus dalam tahap Pengendalian. Berikut 5 standar dengan PTK tertinggi:')

# Sort by PTK descending
ptk_sorted = sorted(STANDAR, key=lambda x: x['ptk'], reverse=True)[:5]
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['No', 'Nama Standar', 'Skor', 'Status', 'PTK'])
for s in ptk_sorted:
    add_table_row(table, [s['no'], s['nama'], s['skor'], s['status'], s['ptk']], align='center')

doc.add_paragraph()

add_heading('5.3 Identifikasi Standar dengan Capaian Terbaik', level=2)
add_body('Standar dengan capaian terbaik (skor tertinggi) dapat menjadi best practice yang dapat direplikasi pada standar lain. Berikut 5 standar dengan skor tertinggi:')

skor_sorted = sorted(STANDAR, key=lambda x: x['skor'], reverse=True)[:5]
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['No', 'Nama Standar', 'Skor', 'Status', 'TL (%)'])
for s in skor_sorted:
    add_table_row(table, [s['no'], s['nama'], s['skor'], s['status'], f"{s['tindakLanjut']}%"], align='center')

doc.add_page_break()

# ============ BAB VI KESIMPULAN ============
add_heading('BAB VI\nKESIMPULAN DAN REKOMENDASI', level=1, color=NAVY, size=20)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading('6.1 Kesimpulan', level=2)
add_body(f'Berdasarkan hasil Assessment AMI TA 2024/2025 terhadap 33 Standar SPMI Universitas Tulungagung, dapat disimpulkan bahwa:')

add_bullet(f'Rata-rata skor AMI institusi mencapai {RATA_SKOR:.1f}/100 yang masuk kategori "Mencapai Standar" (MS). Capaian ini menunjukkan bahwa sistem penjaminan mutu internal Universitas Tulungagung telah berjalan efektif.')
add_bullet(f'Sebanyak {STATUS_COUNTS.get("MS", 0) + STATUS_COUNTS.get("ML", 0)} dari 33 standar ({(STATUS_COUNTS.get("MS", 0) + STATUS_COUNTS.get("ML", 0))/33*100:.1f}%) telah Mencapai atau Melampaui Standar, sementara {STATUS_COUNTS.get("BS", 0)} standar masih Belum mencapai Standar dan memerlukan tindak lanjut.')
add_bullet(f'Total {TOTAL_TEMUAN} temuan dengan {TOTAL_PTK} PTK teridentifikasi, dengan rata-rata tindak lanjut {RATA_TL:.1f}% yang menunjukkan komitmen tinggi dalam penyelesaian temuan.')
add_bullet('Standar Pendidikan dan Penelitian menunjukkan capaian terbaik dengan mayoritas standar MS/ML, sementara Standar Tambahan (khususnya Kerjasama) memerlukan penguatan.')
add_bullet('Tidak ada standar yang menyimpang dari standar (MSV), yang berarti seluruh standar telah berjalan sesuai koridor yang ditetapkan.')

add_heading('6.2 Rekomendasi Tindak Lanjut Prioritas', level=2)
add_body('Berdasarkan analisis temuan, berikut rekomendasi tindak lanjut prioritas untuk tahap Pengendalian dan Peningkatan siklus PPEPP:')

# Tabel rekomendasi
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['Prioritas', 'Standar', 'Isu', 'Rekomendasi'])

rekomendasi = [
    ('1', 'Standar Kerjasama (KJS)', 'Skor 80 (BS), 8 temuan, 5 PTK', 'Menyusun AMI khusus Kerjasama, memperluas MoU/MoA dengan DUDI, meningkatkan implementasi kerjasama Tridharma'),
    ('2', 'Standar Dosen & Tendik (SDT)', 'Sertifikasi pendidik 94,6% (belum 100%)', 'Mempercepat sertifikasi dosen, peningkatan kualifikasi S2/S3 tendik'),
    ('3', 'Standar Isi (SIS)', 'OBE belum konsisten di semua prodi', 'Pelatihan OBE menyeluruh, konsistensi rubrik penilaian berbasis OBE'),
    ('4', 'Standar Penilaian (SPN)', 'Rubrik OBE belum seragam', 'Standardisasi rubrik penilaian, pelatihan teknis untuk dosen'),
    ('5', 'Standar Pengelolaan (SPP)', 'Dashboard SPMI perlu pengembangan', 'Pengembangan fitur dashboard SPMI, integrasi dengan SIAKAD'),
]
for prio, standar, issu, rek in rekomendasi:
    add_table_row(table, [prio, standar, issu, rek], align='left')

doc.add_paragraph()

add_heading('6.3 Penutup', level=2)
add_body('Laporan Hasil Pelaksanaan dan Evaluasi Assessment AMI TA 2024/2025 ini menjadi dasar bagi pimpinan Universitas Tulungagung untuk mengambil kebijakan strategis dalam tahap Pengendalian (P) dan Peningkatan (P) siklus PPEPP. Hasil assessment menunjukkan bahwa sistem penjaminan mutu internal Universitas Tulungagung telah berjalan dengan baik dengan rata-rata skor di atas standar yang ditetapkan. Tindak lanjut terhadap temuan dan PTK akan dilaksanakan melalui Rencana Tindak Lanjut (RTL) yang dibahas dalam Rapat Tinjauan Manajemen (RTM).')

add_body('Pusat Penjaminan Mutu Universitas Tulungagung berkomitmen untuk terus meningkatkan kualitas penyelenggaraan Tridharma Perguruan Tinggi melalui siklus PPEPP yang berkelanjutan, sejalan dengan amanat Permen 39/2025 dan Renstra Universitas Tulungagung 2025-2030.')

doc.add_page_break()

# ============ LAMPIRAN ============
add_heading('LAMPIRAN\nDAFTAR DOKUMEN AMI 2024-2025', level=1, color=NAVY, size=18)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body('Berikut daftar lengkap 15 dokumen AMI TA 2024/2025 yang menjadi sumber data Assessment ini:')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header(table, ['No', 'Nama Dokumen AMI', 'Kode Dokumen', 'Tanggal Audit'])

for i, d in enumerate(DOKUMEN, 1):
    add_table_row(table, [i, d['nama'], d['kode'], d['tanggal']], align='left')

doc.add_paragraph()

add_body('Seluruh dokumen AMI tersebut tersimpan di Pusat Penjaminan Mutu Universitas Tulungagung dan dapat diakses melalui Dashboard SPMI UNITA di https://spmiunita.pages.dev', bold=True, indent=False)

# ============ SAVE ============
output_path = '/home/z/my-project/download/LAPORAN_ASSESSMENT_AMI_2024-2025_UNITA.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"\n✓ Dokumen berhasil disimpan: {output_path}")
print(f"  Ukuran: {os.path.getsize(output_path):,} bytes")

# Also convert to PDF using libreoffice
import subprocess
try:
    result = subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', '/home/z/my-project/download/',
        output_path
    ], capture_output=True, text=True, timeout=120)
    pdf_path = output_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        print(f"  PDF: {pdf_path}")
        print(f"  PDF size: {os.path.getsize(pdf_path):,} bytes")
except Exception as e:
    print(f"  PDF conversion failed: {e}")
